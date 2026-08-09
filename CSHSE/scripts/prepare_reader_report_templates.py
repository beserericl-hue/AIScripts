#!/usr/bin/env python3
"""
Prepare CSHSE Reader Report templates for runtime fill (docx.patchDocument).

The official July-2025 reader report templates (one per degree level) are
compliance checklists: per-standard rows with Compliant / Non-Compliant columns
and a Reader's Comments column. This script inserts {{placeholder}} tokens into
the numbered-standard rows + the Institution/Program header fields, producing
templated copies the Node server fills at runtime (docx.patchDocument).

Run locally with the ai-service venv (has python-docx):
  CSHSE/ai-service/.venv/bin/python scripts/prepare_reader_report_templates.py

Fillable rows are detected in document order (a row with a "Reader's Comments"
cell). The ORDERED `keys` list below assigns each detected row to a standard
number, or None to leave it blank (intro / conditional rows the reader fills).
"""
import re, sys
import docx
from docx.table import Table
from docx.oxml.ns import qn

SRC = "docs"
OUT = "server/src/assets/reader-report-templates"

# Ordered standard-key per detected fillable row, per degree level.
# None = leave blank for the reader (intro / not-applicable rows).
LEVELS = {
    "baccalaureate": {
        "file": "BACCALAUREATE Self_Study_Reader Report template_ Baccalaureate degree  July 2025.docx",
        "keys": ["introduction",None,None,None,None,None,"1","2","3","4","5","6","7","8","9","10",
                 "11","12","13","14","15","16","17","18","19","20","22"],
    },
    "associate": {
        "file": "ASSOCIATE Self_Study_Reader Report template_ Associate degree  July 2025.docx",
        "keys": ["introduction","1","2","3","4","5","6","7","8","9","10",
                 "11","12","13","14","15","16","17","18","19","20"],
    },
    "masters": {
        "file": "Self_Study_Reader_ template_ Master's degree July 2025.docx",
        # 20 fillable header rows in document order: intro, Standards 1-10,
        # curriculum-matrix intro, Standards 11-18.
        "keys": ["introduction","1","2","3","4","5","6","7","8","9","10",
                 None,"11","12","13","14","15","16","17","18"],
        # Masters header is a SINGLE tab-delimited paragraph carrying several
        # labels. docx.patchDocument replaces {{tokens}} INLINE (verified), so we
        # insert them after their labels without disturbing the other fields.
        "multi_label_header": True,
    },
}

def is_fillable(row):
    for c in row.cells:
        t = c.text.lower()
        if "reader" in t and "comment" in t:
            return True
        # Masters template uses an "Areas of noncompliance" comments column.
        if "areas of noncompliance" in t:
            return True
    return False

def find_cells(row):
    """Return (compliant_cell, noncompliant_cell, comments_cell) or (None,None,None)."""
    comp = noncomp = comm = None
    seen = set()
    for c in row.cells:
        if id(c._tc) in seen:  # skip merged duplicates
            continue
        seen.add(id(c._tc))
        t = c.text.strip()
        tl = t.lower()
        norm = re.sub(r"[^a-z]", "", tl)
        if "reader" in tl and "comment" in tl:
            comm = c
        elif "areas of noncompliance" in tl:   # masters comments column
            comm = c
        elif norm == "compliant" or (norm.endswith("compliant")
                                     and "non" not in norm and len(norm) <= 12):
            comp = c   # matches "Compliant" and masters' "a.  Compliant"
        elif norm.startswith("noncompliant"):
            noncomp = c
    return comp, noncomp, comm

def prepend_run(cell, text):
    p = cell.paragraphs[0]
    r = p.add_run(text)
    # move the new run to the front of the paragraph
    p._p.insert(0, r._r)

def fill_header(doc, label, token):
    """Replace the underscores after a 'Label:' paragraph with a token."""
    for p in doc.paragraphs:
        t = p.text
        if t.strip().startswith(label):
            # keep label, drop underscores, add token
            for r in list(p.runs):
                r._r.getparent().remove(r._r)
            p.add_run(f"{label} {token}")
            return True
    return False

def fill_multi_label_header(doc):
    """Masters header is ONE paragraph with several tab-separated labels
    ("Institution's Name: \\t Program's Name: \\t Program Director's Name \\t
    Reader's Name ... Date:"). Insert {{inst_name}} / {{prog_name}} INLINE right
    after their labels — patchDocument replaces tokens in place, so the other
    labels are preserved. Handles straight (') and curly (’) apostrophes."""
    inst = prog = False
    for p in doc.paragraphs:
        t = p.text
        if ("Institution" in t and "Program" in t and "Name" in t
                and "{{inst_name}}" not in t):
            new, n1 = re.subn(r"(Institution[’']s Name:)", r"\1 {{inst_name}}", t, count=1)
            new, n2 = re.subn(r"(Program[’']s Name:)", r"\1 {{prog_name}}", new, count=1)
            inst, prog = bool(n1), bool(n2)
            # rebuild as a single run, preserving the first run's font/weight
            first = p.runs[0] if p.runs else None
            for r in list(p.runs):
                r._r.getparent().remove(r._r)
            run = p.add_run(new)
            if first is not None:
                run.bold = first.bold
                run.italic = first.italic
                if first.font.size: run.font.size = first.font.size
                if first.font.name: run.font.name = first.font.name
            return inst, prog
    return inst, prog

def tables(doc):
    for blk in doc.element.body.iterchildren():
        if blk.tag == qn("w:tbl"):
            yield Table(blk, doc)

def process(level, cfg):
    path = f"{SRC}/{cfg['file']}"
    doc = docx.Document(path)
    # fillable rows in document order
    rows = []
    for tbl in tables(doc):
        for row in tbl.rows:
            if is_fillable(row):
                rows.append(row)
    keys = cfg["keys"]
    if len(rows) != len(keys):
        print(f"  !! {level}: detected {len(rows)} fillable rows but keys has {len(keys)} — ABORT")
        return False
    filled = 0
    for row, key in zip(rows, keys):
        if key is None:
            continue
        comp, noncomp, comm = find_cells(row)
        if comp is not None:  prepend_run(comp, f"{{{{c_{key}}}}} ")
        if noncomp is not None: prepend_run(noncomp, f"{{{{n_{key}}}}} ")
        if comm is not None:  comm.add_paragraph(f"{{{{cm_{key}}}}}")
        filled += 1
    if cfg.get("multi_label_header"):
        ih, ph = fill_multi_label_header(doc)
    elif cfg.get("no_header"):
        ih = ph = False
    else:
        ih = fill_header(doc, "Institution’s Name:", "{{inst_name}}")
        ph = fill_header(doc, "Program’s Name:", "{{prog_name}}")
    out = f"{OUT}/{level}.docx"
    doc.save(out)
    print(f"  {level}: {filled} standard rows templated, header(inst={ih},prog={ph}) -> {out}")
    return True

ok = True
for level, cfg in LEVELS.items():
    print(f"== {level} ==")
    ok = process(level, cfg) and ok
sys.exit(0 if ok else 1)
