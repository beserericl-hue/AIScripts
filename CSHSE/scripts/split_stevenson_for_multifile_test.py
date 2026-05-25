#!/usr/bin/env python3
"""Throwaway: split a CSHSE self-study DOCX into per-section files for
the multi-file wizard test.

NOT delivered in the system. Run locally to produce a small folder of
related .docx files we can drag-and-drop onto the wizard's multi-file
Upload step (CR-041) to exercise the batched-import flow.

Usage:
    python3 split_stevenson_for_multifile_test.py /path/to/stevenson.docx

Output: writes new files into the same directory as the input. Existing
output files are overwritten. Files produced:

    stevenson__01-standards-01-05.docx     (Standard 1 → Standard 5)
    stevenson__02-standards-06-09.docx     (Standard 6 → Standard 9)
    stevenson__03-standards-10-13.docx     (Standard 10 → Standard 13)
    stevenson__04-standards-14-21.docx     (Standard 14 → Standard 21)
    stevenson__05-appendix.docx            (everything after the last Standard)

Plus, when the original contains identifiable blocks, optional standalones:
    stevenson__cv-only__<faculty>.docx   (faculty CV blocks, if detected)
    stevenson__paper__<title>.docx       (research paper / report blocks)
    stevenson__syllabus__<code>.docx     (syllabus blocks)

The chunker is intentionally rough — it splits on \"Standard N\" headings
and uses cheap regex for the appendix items. False positives become
extra files the coordinator can drop or skip in the wizard.

Dependencies: python-docx (already in ai-service/requirements.txt; install
locally with `pip install python-docx` if running outside the venv).
"""
from __future__ import annotations

import os
import re
import sys
from pathlib import Path
from typing import Iterable

try:
    from docx import Document  # type: ignore[import-untyped]
    from docx.document import Document as _DocumentT  # type: ignore[import-untyped]
except ImportError:
    sys.stderr.write(
        "python-docx is required. Install with: pip install python-docx\n"
    )
    sys.exit(1)


# Headings that delimit a Standard. Matches "Standard 1", "STANDARD 12.a", etc.
_STANDARD_RE = re.compile(r"^\s*STANDARD\s+(\d{1,2})\b", re.IGNORECASE)
# Appendix items inside Stevenson's tail. These are best-effort regex.
_APPENDIX_PAPER_RE = re.compile(
    r"(sample\s+\w+\s+(report|paper|project|essay|reflection|interview)"
    r"|^(.*\(\d+\s*points\))\s*$)",
    re.IGNORECASE,
)
_APPENDIX_SYLLABUS_RE = re.compile(
    r"(course\s+syllabus|syllabus|^[A-Z]{2,5}\s+\d{2,4}\b)",
    re.IGNORECASE,
)
# Faculty CV: name line + "EDUCATION" within ~10 paragraphs.
_FACULTY_NAME_RE = re.compile(
    r"^\s*((?:Dr\.\s+)?[A-Z][a-z]+\s+(?:[A-Z]\.\s+)?[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\s*$"
)
_EDUCATION_RE = re.compile(r"^\s*EDUCATION\b", re.IGNORECASE)


def _slugify(s: str, max_len: int = 50) -> str:
    s = re.sub(r"[^A-Za-z0-9]+", "-", s).strip("-").lower()
    return s[:max_len] or "untitled"


def _paragraphs_with_styles(doc: _DocumentT) -> list[tuple[str, str]]:
    """Return (style_name, text) for every paragraph."""
    return [(p.style.name if p.style else "Normal", p.text) for p in doc.paragraphs]


def _split_by_standards(paragraphs: list[tuple[str, str]]) -> dict[int, list[int]]:
    """Return {standard_number: [paragraph_index, paragraph_index, ...]}.

    Standard 0 is the preamble (everything before Standard 1).
    Standard 99 is the appendix (everything after the last Standard).
    """
    current = 0
    saw_any_standard = False
    by_standard: dict[int, list[int]] = {0: []}
    for idx, (_style, text) in enumerate(paragraphs):
        m = _STANDARD_RE.match(text)
        if m and not saw_any_standard:
            # First Standard found — open it.
            current = int(m.group(1))
            saw_any_standard = True
            by_standard.setdefault(current, []).append(idx)
            continue
        if m and saw_any_standard:
            new_std = int(m.group(1))
            if new_std != current:
                current = new_std
                by_standard.setdefault(current, []).append(idx)
                continue
        by_standard.setdefault(current, []).append(idx)
    return by_standard


def _write_docx(
    src: _DocumentT,
    paragraph_indices: Iterable[int],
    out_path: Path,
    title: str,
) -> None:
    """Create a new .docx containing only the selected paragraphs.

    python-docx doesn't have a clean "copy paragraph" API; we recreate
    each paragraph by reading its text + applying the source style name.
    Images and complex formatting are lost — that's fine for the
    multi-file test fixture.
    """
    out = Document()
    out.add_paragraph(title, style="Heading 1")
    src_paragraphs = list(src.paragraphs)
    for i in paragraph_indices:
        if i < 0 or i >= len(src_paragraphs):
            continue
        sp = src_paragraphs[i]
        style_name = sp.style.name if sp.style else "Normal"
        try:
            out.add_paragraph(sp.text, style=style_name)
        except KeyError:
            # Style not found in the new doc's catalog; fall back to Normal.
            out.add_paragraph(sp.text)
    out.save(out_path)


def _is_full_cv_anchor(paragraphs: list[tuple[str, str]], idx: int) -> str | None:
    """Return the captured name if `paragraphs[idx]` is a real CV anchor
    (name-shaped line followed by EDUCATION within 10 paragraphs), else
    None. Centralises the dual-signal check so block-extension below
    can reuse it and not mis-fire on institution names like ``Loyola
    University Maryland`` that happen to be 3 capitalised tokens.
    """
    if idx < 0 or idx >= len(paragraphs):
        return None
    nm = _FACULTY_NAME_RE.match(paragraphs[idx][1])
    if not nm:
        return None
    for j in range(idx + 1, min(idx + 11, len(paragraphs))):
        if _EDUCATION_RE.match(paragraphs[j][1]):
            return nm.group(1).strip()
    return None


def _bookmark_cv_blocks(doc: _DocumentT) -> list[tuple[str, list[int]]] | None:
    """If the source document carries ``FacCVs*`` bookmarks (Stevenson
    uses these to anchor the TOC entries that point to each faculty
    CV), use them as authoritative CV boundaries. Returns
    [(faculty_name, [paragraph_indices])] in document order, or None if
    no such bookmarks exist (caller falls back to heuristic detection).

    Faculty names are pulled from the paragraph nearest the bookmark
    (some bookmarks attach to empty paragraphs that precede the actual
    CV body) so the produced filenames are coordinator-friendly
    (``barry-w-thomas`` rather than the raw bookmark slug ``Thomas``).
    """
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    para_elements = [p._element for p in doc.paragraphs]
    para_idx_by_element = {id(el): i for i, el in enumerate(para_elements)}

    bookmark_positions: list[tuple[int, str]] = []
    for el in doc.element.iter():
        if el.tag != f"{W}bookmarkStart":
            continue
        name = el.attrib.get(f"{W}name", "")
        if not name.startswith("FacCVs"):
            continue
        parent = el
        while parent is not None and parent.tag != f"{W}p":
            parent = parent.getparent() if hasattr(parent, "getparent") else None
        if parent is None or id(parent) not in para_idx_by_element:
            continue
        para_idx = para_idx_by_element[id(parent)]
        if name == "FacCVs":
            # Parent "Faculty Curriculum Vitae" anchor — skip; the
            # per-faculty bookmarks (FacCVsRosicky, FacCVsSwish, ...)
            # carry the precise per-CV starts.
            continue
        bookmark_positions.append((para_idx, name))

    if not bookmark_positions:
        return None

    # Dedupe duplicate bookmarks at the same paragraph index (Stevenson
    # has both FacCVsLess + FacCVsLesser pointing at Loryn Lesser).
    seen_idx: set[int] = set()
    deduped: list[tuple[int, str]] = []
    for idx, name in sorted(bookmark_positions):
        if idx in seen_idx:
            continue
        seen_idx.add(idx)
        deduped.append((idx, name))
    bookmark_positions = deduped

    total_paragraphs = len(doc.paragraphs)
    out: list[tuple[str, list[int]]] = []
    for i, (start_idx, raw_name) in enumerate(bookmark_positions):
        if i + 1 < len(bookmark_positions):
            end_idx = bookmark_positions[i + 1][0]
        else:
            # Last CV — cap at the next "Standard N" heading if one
            # appears within 250 paragraphs, else absolute cap.
            end_idx = min(start_idx + 250, total_paragraphs)
            for k in range(start_idx + 5, end_idx):
                text = doc.paragraphs[k].text or ""
                if _STANDARD_RE.match(text):
                    end_idx = k
                    break
        # Resolve a coordinator-friendly faculty name. Some bookmarks
        # land on empty paragraphs (page breaks / separators) — scan
        # forward up to 5 paragraphs for the first non-empty line that
        # matches the name regex.
        faculty_name: str | None = None
        for k in range(start_idx, min(start_idx + 5, end_idx)):
            text = (doc.paragraphs[k].text or "").strip()
            if not text:
                continue
            # Strip honorific / credential suffix (e.g. ", J.D., HS-BCP")
            head = re.split(r"[,(]", text, maxsplit=1)[0].strip()
            head = re.sub(r"^(Dr|Prof|Mr|Mrs|Ms|Mx)\.?\s+", "", head, flags=re.IGNORECASE)
            tokens = head.split()
            if 2 <= len(tokens) <= 5 and all(t[:1].isalpha() for t in tokens):
                faculty_name = head
                break
        if not faculty_name:
            faculty_name = raw_name[len("FacCVs"):] or f"cv-{start_idx}"
        out.append((faculty_name, list(range(start_idx, end_idx))))
    return out


def _detect_cv_blocks(paragraphs: list[tuple[str, str]]) -> list[tuple[str, list[int]]]:
    """Heuristic fallback for documents without FacCVs* bookmarks.

    Name-shaped line + EDUCATION within 10 paras. Returns
    [(faculty_name, [paragraph_indices])]. Each block extends until
    the next *real* CV anchor (also name + EDUCATION) or until 240
    paragraphs pass — long enough to contain a full Stevenson CV
    (which routinely runs 150+ paragraphs once Education + Academic
    Employment + Teaching Experiences + Publications + Service are
    all included).
    """
    out: list[tuple[str, list[int]]] = []
    i = 0
    while i < len(paragraphs):
        name = _is_full_cv_anchor(paragraphs, i)
        if not name:
            i += 1
            continue
        # Extend block until the next *real* anchor (name + EDUCATION
        # within 10 paras) or up to 240 paragraphs. Plain name-shaped
        # lines (e.g. institution names like "Loyola University Maryland")
        # do NOT close the block — they were the source of the
        # previous truncation bug that lopped Barry W. Thomas's CV off
        # at the first institution name inside his Education block.
        end = min(i + 240, len(paragraphs))
        for k in range(i + 6, end):
            if _is_full_cv_anchor(paragraphs, k):
                end = k
                break
        out.append((name, list(range(i, end))))
        i = end
    return out


def _detect_paper_blocks(paragraphs: list[tuple[str, str]]) -> list[tuple[str, list[int]]]:
    out: list[tuple[str, list[int]]] = []
    i = 0
    while i < len(paragraphs):
        _, text = paragraphs[i]
        if _APPENDIX_PAPER_RE.search(text):
            title = text.strip()[:40] or f"paper-{i}"
            end = min(i + 80, len(paragraphs))
            # Stop early at the next paper / syllabus / standard.
            for k in range(i + 1, end):
                t = paragraphs[k][1]
                if (
                    _APPENDIX_PAPER_RE.search(t)
                    or _APPENDIX_SYLLABUS_RE.search(t)
                    or _STANDARD_RE.match(t)
                ):
                    end = k
                    break
            out.append((title, list(range(i, end))))
            i = end
        else:
            i += 1
    return out


def _detect_syllabus_blocks(paragraphs: list[tuple[str, str]]) -> list[tuple[str, list[int]]]:
    out: list[tuple[str, list[int]]] = []
    i = 0
    seen_titles: set[str] = set()
    while i < len(paragraphs):
        _, text = paragraphs[i]
        if _APPENDIX_SYLLABUS_RE.search(text):
            title = text.strip()[:40] or f"syllabus-{i}"
            if title in seen_titles:
                i += 1
                continue
            seen_titles.add(title)
            end = min(i + 60, len(paragraphs))
            for k in range(i + 1, end):
                t = paragraphs[k][1]
                if (
                    _APPENDIX_SYLLABUS_RE.search(t)
                    or _APPENDIX_PAPER_RE.search(t)
                    or _STANDARD_RE.match(t)
                ):
                    end = k
                    break
            out.append((title, list(range(i, end))))
            i = end
        else:
            i += 1
    return out


def main(argv: list[str]) -> int:
    if len(argv) != 2:
        sys.stderr.write(
            "Usage: split_stevenson_for_multifile_test.py /path/to/stevenson.docx\n"
        )
        return 2
    src_path = Path(argv[1]).expanduser().resolve()
    if not src_path.exists():
        sys.stderr.write(f"Input not found: {src_path}\n")
        return 1
    if src_path.suffix.lower() != ".docx":
        sys.stderr.write(f"Expected a .docx file, got: {src_path.suffix}\n")
        return 1

    out_dir = src_path.parent
    base = src_path.stem
    print(f"Reading {src_path}")
    src = Document(str(src_path))
    paragraphs = _paragraphs_with_styles(src)
    print(f"  {len(paragraphs)} paragraphs")

    by_standard = _split_by_standards(paragraphs)
    print(f"  found Standards: {sorted(k for k in by_standard if 1 <= k <= 21)}")

    # Group standards into 4 chunks for the multi-file test.
    groups = [
        ("01-standards-01-05", range(1, 6)),
        ("02-standards-06-09", range(6, 10)),
        ("03-standards-10-13", range(10, 14)),
        ("04-standards-14-21", range(14, 22)),
    ]
    for label, stds in groups:
        indices: list[int] = []
        for std in stds:
            indices.extend(by_standard.get(std, []))
        if not indices:
            print(f"  skip {label}: no paragraphs")
            continue
        out_path = out_dir / f"{base}__{label}.docx"
        _write_docx(src, indices, out_path, f"{base} — {label.replace('-', ' ').title()}")
        print(f"  wrote {out_path.name} ({len(indices)} paragraphs)")

    # Preamble + appendix as separate files.
    if by_standard.get(0):
        out_path = out_dir / f"{base}__00-preamble.docx"
        _write_docx(src, by_standard[0], out_path, f"{base} — preamble")
        print(f"  wrote {out_path.name} ({len(by_standard[0])} paragraphs)")

    appendix_paragraphs: list[int] = []
    last_std = max((k for k in by_standard if 1 <= k <= 21), default=0)
    if last_std:
        for std, idxs in by_standard.items():
            if std > last_std:
                appendix_paragraphs.extend(idxs)
    # Anything after the last Standard's paragraph index counts as appendix.
    if last_std and by_standard.get(last_std):
        last_paragraph = max(by_standard[last_std])
        appendix_paragraphs.extend(
            i for i in range(last_paragraph + 1, len(paragraphs))
        )
    appendix_paragraphs = sorted(set(appendix_paragraphs))
    if appendix_paragraphs:
        out_path = out_dir / f"{base}__05-appendix.docx"
        _write_docx(src, appendix_paragraphs, out_path, f"{base} — appendix")
        print(f"  wrote {out_path.name} ({len(appendix_paragraphs)} paragraphs)")

    # Optional standalone CV / paper / syllabus files (for CR-033 + CR-040
    # standalone-upload testing). Prefer FacCVs* bookmarks when present —
    # Stevenson's source doc has them and they give us byte-exact CV
    # boundaries; the heuristic fallback over-extends past the actual CV
    # end-of-content (it ran 240 paragraphs forward and swallowed the
    # "Academic Affairs Committee" section that follows the last CV).
    bookmark_cvs = _bookmark_cv_blocks(src)
    if bookmark_cvs is not None:
        cvs = bookmark_cvs
    else:
        cvs = _detect_cv_blocks(paragraphs)
    for name, idxs in cvs:  # bookmark path is authoritative — emit all
        slug = _slugify(name)
        out_path = out_dir / f"{base}__cv-only__{slug}.docx"
        _write_docx(src, idxs, out_path, f"CV — {name}")
        print(f"  wrote {out_path.name} (CV; {len(idxs)} paragraphs)")
    papers = _detect_paper_blocks(paragraphs)
    for title, idxs in papers[:5]:
        slug = _slugify(title)
        out_path = out_dir / f"{base}__paper__{slug}.docx"
        _write_docx(src, idxs, out_path, f"Paper — {title}")
        print(f"  wrote {out_path.name} (paper; {len(idxs)} paragraphs)")
    syllabi = _detect_syllabus_blocks(paragraphs)
    for title, idxs in syllabi[:5]:
        slug = _slugify(title)
        out_path = out_dir / f"{base}__syllabus__{slug}.docx"
        _write_docx(src, idxs, out_path, f"Syllabus — {title}")
        print(f"  wrote {out_path.name} (syllabus; {len(idxs)} paragraphs)")

    print("Done. Drag any combination of these files onto the wizard's")
    print("Upload step to exercise the multi-file batched-import flow.")
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv))
