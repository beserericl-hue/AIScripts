"""Walker for the CSHSE Self-Study Template DOCX format.

The template format is the spec-as-outline alternative to a free-form
self-study like Stevenson's: each section is named directly after a
prompt from the Handbook (``1. Specify the degree(s)…``, ``2a. Describe
the organizational structure…``, ``Standard 1, Specification a.``) and
the institution writes a ``Response:`` underneath each. This format is
what institutions just starting accreditation receive as a writing
template, and they re-import it repeatedly as more sections are filled.

What this walker does:

  1. Reads the DOCX directly with python-docx (no HTML conversion).
  2. Walks paragraphs in document order. A paragraph whose text matches
     one of the template heading patterns starts a new section. Body
     paragraphs accumulate to the most recent heading.
  3. Treats ``Response:`` as a marker within a section, not a section
     boundary. The body after it is the institution's content.
  4. Treats ``See Appendix …`` lines as content (the matcher uses them
     as evidence pointers, not as cuts).
  5. Detects empty / ``Not applicable`` placeholder responses so the
     preview can show them as "section started but unwritten".
  6. Emits ``Section`` objects shaped exactly like the other walkers so
     the matcher + bucket allocator + coverage reviewer all run
     unchanged downstream.

What this walker does NOT do:

  - Table extraction. The template format embeds the Table of Contents,
    a Curriculum Requirements credit-hours table, and an Appendix list
    as tables; none of those carry per-spec narrative content. The
    matrix data extractor (``app/matrix/data_extractor.py``) handles
    matrix DOCX/tables separately when present.
  - HTML conversion. The template format ships as DOCX and is small
    (~400 KB) so we read paragraphs directly. A Stevenson-style HTML
    pipeline would be overkill.
"""
from __future__ import annotations

import base64
import re
import uuid
from dataclasses import dataclass, field

from docx import Document

# OOXML namespaces for the rich-HTML walk (hyperlinks + inline images).
_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_R = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"
_A = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
# Cap an inline (base64) image so a doc full of large scans can't blow the
# import callback payload. Bigger images are noted and left to be attached via
# the Supporting File Library.
# CR-063 — inline images up to this size as base64 data URIs. Raised from
# 1.2 MB → 2 MB so realistic embedded images (maps, charts, scanned letters)
# are never dropped. Above this, base64 would risk the 16 MB Mongo doc limit on
# aiReviewState, so a clear placeholder points the PC to the Supporting File
# Library instead (a true never-drop route-store is a server-side follow-on).
_MAX_INLINE_IMAGE_BYTES = 2_000_000

from app.splitter.sections import Section, _heuristic_flags

# A "See Appendix …" / "See Part II Appendices …" reference. In real PC
# submissions these are REMINDERS for the coordinator to import the named
# appendix document as supporting evidence — never extractable content. We
# surface them as import reminders, we do not treat them as evidence files.
_APPENDIX_RE = re.compile(
    r"see\s+(?:part\s+[ivx]+\s+)?append(?:ix|ices)\b[^\n]{0,160}",
    re.IGNORECASE,
)

# ----------------------------------------------------------------- patterns
#
# Ordered from most-specific to most-general. The first pattern to match a
# paragraph's text wins. Patterns capture standard / spec hints when they
# appear in the heading itself; the matcher uses these as doc-letter /
# doc-standard-hint signals.

_HEADING_PATTERNS = [
    # "Standard 1, Specification a." or "Standard 11 Specification b"
    re.compile(
        r"^\s*Standard\s+(?P<std>\d{1,2})\s*[,.\s]+\s*Specification\s+(?P<spec>[a-j])\b",
        re.IGNORECASE,
    ),
    # "Standard 12b." / "Standard 16c"
    re.compile(
        r"^\s*Standard\s+(?P<std>\d{1,2})\s*(?P<spec>[a-j])\.?\s+",
        re.IGNORECASE,
    ),
    # "Standard 1:" / "Standard 4: The program shall…" — the colon form a
    # program uses to open a Standard whose normative ("shall") sentence runs
    # ON THE SAME LINE. It's a Standard ROOT: a std hint, NO spec. The specs
    # (1a., 1b., …) follow underneath. (Must precede the bare-standard rule.)
    re.compile(r"^\s*Standard\s+(?P<std>\d{1,2})\s*:", re.IGNORECASE),
    # "1.a." / "11.b" — dotted spec form (Handbook canonical form)
    re.compile(r"^\s*(?P<std>\d{1,2})\.(?P<spec>[a-j])\.?\s"),
    # "1a." / "2b" — undotted spec form (template intro-section form)
    re.compile(r"^\s*(?P<std>\d{1,2})(?P<spec>[a-j])\.?\s"),
    # Section root like "1." / "2." / "Standard 1" — std hint without spec
    re.compile(r"^\s*(?P<std>\d{1,2})\.\s"),
    re.compile(r"^\s*Standard\s+(?P<std>\d{1,2})\b\s*$", re.IGNORECASE),
    # "A. Required Self-Study Introductory…", "B. Glossary", etc.
    re.compile(r"^\s*(?P<letter>[A-Z])\.\s+[A-Z]"),
]

# The heading that OPENS the standards region. Everything BEFORE the first one
# of these is the document Introduction (front matter + the intro's own
# "1." / "2a." prompts + the closing Glossary of terms). The user's rule:
# "document introduction ends with glossary of terms" — and the glossary always
# sits just before the first Standard.
#
# Match ONLY a true boundary: a Roman-numeral part header ("I. GENERAL PROGRAM
# CHARACTERISTICS") or a Standard ROOT — "Standard 1:" (colon) or a bare
# "Standard 1" on its own line. Deliberately NOT inline curriculum citations
# like "Standard 13 a:" or "Standard 12b. …", which the front-matter program
# description sprinkles BEFORE the glossary — those must not end the intro early
# (that was dropping the glossary out of the introduction).
#
# Why the region matters: the Introduction reuses spec-LOOKING numbering ("2a.
# Describe the organizational structure" is an INTRO prompt, not Standard 2 spec
# a). Held inside the intro region, those hints are dropped so nothing misfiles
# under a Standard — the whole region routes to introduction:document instead.
_STANDARDS_REGION_RE = re.compile(
    r"^\s*(?:[IVX]+\.\s+[A-Z]|Standard\s+\d{1,2}\s*(?::|\.?\s*$))",
    re.IGNORECASE,
)

# Markers that are not section breaks — they appear inside a section body.
_RESPONSE_MARKER_RE = re.compile(r"^\s*Response\s*:\s*", re.IGNORECASE)

# Placeholder responses the institution hasn't filled in.
_PLACEHOLDER_RE = re.compile(
    r"^\s*(not\s+applicable|n/?a|tbd|to\s+be\s+determined|\[response\]|insert\s+response)\s*\.?\s*$",
    re.IGNORECASE,
)

# Headings that are document-front matter (skip — they don't map to specs).
_FRONT_MATTER_HEADINGS = re.compile(
    r"^\s*(table of contents|introduction and instructions|"
    r"council for standards in human service education|self-study template)\b",
    re.IGNORECASE,
)


# ---------------------------------------------------------------- public types


@dataclass
class TemplateSection:
    """A heading + accumulated body from the template DOCX, pre-Section.

    Kept as a thin intermediate so callers (and tests) can inspect what
    the walker found before it's wrapped into the downstream ``Section``
    dataclass. The ``placeholder`` flag is True when the response body
    is empty or matches the ``_PLACEHOLDER_RE`` patterns ("Not
    applicable", "TBD", etc.) — the wizard surfaces these as
    "started-but-unwritten" rather than passing them to the matcher.
    """
    paragraph_index: int  # 0-based index of the heading paragraph in the DOCX
    heading: str
    body_paragraphs: list[str]
    standard_hint: str | None
    spec_hint: str | None
    placeholder: bool
    # True when this section lives in the document Introduction region (before
    # the first Standard / Roman part header — title block, intro prompts,
    # glossary). These route to introduction:document, and their spec-looking
    # numbering is NOT trusted as Standard hints.
    is_introduction: bool = False
    # Embedded TABLES that fall under this section in document order (faculty
    # rosters, advisory-board attendee lists, credit-hour grids). The paragraph
    # walk drops tables entirely, so these are captured separately and surfaced
    # as SUPPORTING EVIDENCE for the section's spec. Each: {"html","text"}.
    evidence_tables: list[dict] = field(default_factory=list)
    # "See Appendix …" references found in this section's prose — surfaced as
    # PC import reminders (import that appendix doc as supporting evidence).
    appendix_refs: list[str] = field(default_factory=list)
    # Rich-HTML body parts (paragraphs with hyperlinks/images + inline tables)
    # in document order — the faithful narrative content for the editor.
    body_html_parts: list[str] = field(default_factory=list)

    @property
    def body_html(self) -> str:
        html = "".join(self.body_html_parts).strip()
        # CR-062 — wrap each run of consecutive <li> items in a <ul> so Word
        # lists render as lists (with their hyperlinks intact) instead of a
        # flat sequence of paragraphs.
        html = re.sub(r"(?:<li>.*?</li>)+", lambda m: f"<ul>{m.group(0)}</ul>", html, flags=re.DOTALL)
        return html

    @property
    def body_text(self) -> str:
        return "\n\n".join(p for p in self.body_paragraphs if p).strip()

    @property
    def word_count(self) -> int:
        return len(self.body_text.split())


# ---------------------------------------------------------------- helpers


def _is_front_matter(text: str) -> bool:
    """``True`` for template title / instructions / TOC headings — these are
    skipped entirely (neither a section cut nor accumulated as body) so
    the title-page lines under them don't bleed into a pseudo-section.
    """
    return bool(_FRONT_MATTER_HEADINGS.match(text))


def _match_heading(text: str) -> tuple[str | None, str | None] | None:
    """Return ``(std_hint, spec_hint)`` if ``text`` matches a heading pattern.

    The hints may both be None even on a match (e.g. ``"B. Glossary"`` is
    a heading but doesn't carry a numeric standard hint). Front matter
    returns ``(None, None)`` so the walker can detect it and skip — see
    ``_is_front_matter``.
    """
    if _is_front_matter(text):
        return (None, None)
    for pat in _HEADING_PATTERNS:
        m = pat.match(text)
        if not m:
            continue
        std = m.groupdict().get("std")
        spec = m.groupdict().get("spec")
        return (std, spec.lower() if spec else None)
    return None


def _strip_response_marker(text: str) -> str:
    """Drop a leading ``Response:`` marker; return the remainder (possibly empty)."""
    return _RESPONSE_MARKER_RE.sub("", text).strip()


def _extract_appendix_refs(text: str) -> list[str]:
    """Return any ``See Appendix …`` references in ``text`` (verbatim, trimmed)."""
    return [m.group(0).strip() for m in _APPENDIX_RE.finditer(text or "")]


def _esc(s: str) -> str:
    return (s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _table_to_text(table) -> str:
    return "\n".join(
        " | ".join((c.text or "").strip() for c in r.cells) for r in table.rows
    ).strip()


def _table_to_html(table) -> str:
    rows = [
        "<tr>" + "".join(f"<td>{_esc((c.text or '').strip())}</td>" for c in r.cells) + "</tr>"
        for r in table.rows
    ]
    return "<table>" + "".join(rows) + "</table>"


# ---------------------------------------------------------------- tabular template
#
# Many institutions submit the OFFICIAL CSHSE template with its TABLES intact —
# one table per Standard, the "Standard N:" heading + lettered specs (a./b./…) +
# each "Response:" all inside table CELLS. (Kennesaw removed the tables, leaving
# a paragraph template; several schools keep them.) The paragraph walk can't see
# cell content, so we DECOMPOSE each standard/spec response table back into the
# paragraph-block stream the existing walker already understands. Non-standard
# tables (matrix grids, credit-hour tables, rosters) are left as opaque evidence
# tables — exactly as before — so the Kennesaw + MCC paths are unaffected.

_STD_HEADER_CELL_RE = re.compile(r"^\s*Standard\s+(?P<std>\d{1,2})\s*:", re.IGNORECASE)
_SPEC_LETTER_CELL_RE = re.compile(r"^\s*(?P<letter>[a-j])[.)]?\s*$", re.IGNORECASE)
_INTRO_NUM_CELL_RE = re.compile(r"^\s*(?P<num>\d{1,2})[.)]?\s*$")
_GLOSSARY_RE = re.compile(r"glossary", re.IGNORECASE)


def _cell_text(cell) -> str:
    return (cell.text or "").strip()


def _cell_html(cell) -> str:
    """Rich inner HTML for a table cell — joins each paragraph's HTML (keeping
    hyperlinks + inline images), each wrapped in <p>."""
    parts = []
    for p in cell.paragraphs:
        h = _paragraph_to_html(p)
        if h.strip():
            parts.append(f"<p>{h}</p>")
    return "".join(parts)


_RESP_STANDALONE_RE = re.compile(r"^\s*Response\s*:?\.?\s*$", re.IGNORECASE)
_RESP_INLINE_RE = re.compile(r"\bResponse\s*:\s*", re.IGNORECASE)


def _para_body_block(p) -> tuple:
    """A ('p', text, html, is_list) body block for one cell paragraph, keeping
    hyperlinks / inline images / list formatting."""
    text = (p.text or "").strip()
    html = _paragraph_to_html(p)
    is_list = _is_list_item(p._p)
    if html.strip():
        return ("p", text, f"<li>{html}</li>" if is_list else f"<p>{html}</p>", is_list)
    return ("p", text, f"<p>{_esc(text)}</p>", is_list)


def _split_cell(cell) -> tuple[str, list[tuple]]:
    """Split a spec cell into ``(prompt_text, body_blocks)`` where body_blocks are
    the institution's answer paragraphs (as ('p', text, html, is_list) blocks).

    Robust to the three answer shapes seen in real templates:
      * a standalone ``Response:`` / ``Response`` marker paragraph,
      * an inline ``… Response: …`` inside a paragraph, and
      * NO marker at all — the institution wrote the answer directly under the
        Handbook prompt (first paragraph is the prompt, the rest is the answer).
    """
    paras = list(cell.paragraphs)
    # 1) standalone "Response" / "Response:" marker paragraph
    for i, p in enumerate(paras):
        if _RESP_STANDALONE_RE.match(p.text or ""):
            prompt = " ".join((q.text or "").strip() for q in paras[:i] if (q.text or "").strip())
            body = [_para_body_block(q) for q in paras[i + 1:] if (q.text or "").strip()]
            return (prompt.strip(), body)
    # 2) inline "Response:" inside a paragraph
    for i, p in enumerate(paras):
        t = p.text or ""
        m = _RESP_INLINE_RE.search(t)
        if m:
            pre = " ".join((q.text or "").strip() for q in paras[:i] if (q.text or "").strip())
            prompt = (pre + " " + t[: m.start()]).strip()
            after = t[m.end():].strip()
            body: list[tuple] = []
            if after:
                body.append(("p", after, f"<p>{_esc(after)}</p>", False))
            body += [_para_body_block(q) for q in paras[i + 1:] if (q.text or "").strip()]
            return (prompt, body)
    # 3) no marker — first non-empty paragraph is the prompt, the rest the answer
    nonempty = [p for p in paras if (p.text or "").strip()]
    if not nonempty:
        return ("", [])
    prompt = (nonempty[0].text or "").strip()
    body = [_para_body_block(p) for p in nonempty[1:]]
    return (prompt, body)


def _looks_like_standard_table(table) -> bool:
    """True when any cell in the table opens with ``Standard N:`` — i.e. this is a
    per-Standard response table of the tabular official template."""
    try:
        for row in table.rows:
            for cell in row.cells:
                if _STD_HEADER_CELL_RE.match(_cell_text(cell)):
                    return True
    except Exception:  # noqa: BLE001
        pass
    return False


def _first_spec_letter(table) -> str | None:
    for row in table.rows:
        m = _SPEC_LETTER_CELL_RE.match(_cell_text(row.cells[0]))
        if m:
            return m.group("letter").lower()
    return None


def _is_marker_only(text: str) -> bool:
    """True when a cell holds ONLY a spec letter / sub-number / std header — no
    real content. Some tables duplicate the ``a.`` marker across two columns
    (AACC Standard 16 is 3-wide: ``a. | a. | <prompt+response>``)."""
    t = (text or "").strip()
    if not t:
        return True
    return bool(
        _SPEC_LETTER_CELL_RE.match(t)
        or _INTRO_NUM_CELL_RE.match(t)
        or _STD_HEADER_CELL_RE.match(t)
    )


def _content_cell(row):
    """The cell in a row that holds the prompt/response — the LAST cell that
    isn't a bare marker. Handles 2-column (``a. | content``) and 3-column
    (``a. | a. | content``) response tables alike."""
    cells = list(row.cells)
    for c in reversed(cells):
        if not _is_marker_only(_cell_text(c)):
            return c
    return cells[-1]


def _decompose_response_table(table, state: dict) -> list[tuple] | None:
    """Decompose a per-Standard response table into ('p', text, html, is_list)
    blocks the paragraph walker understands. Returns ``None`` when the table is
    NOT a standard/spec response table (glossary, matrix, roster, credit grid) so
    the caller keeps it as an opaque evidence table.

    Handles the two real-world wrinkles:
      * headerless CONTINUATION tables — Standards 5 and 10 in the AACC template
        have no ``Standard N:`` header cell. A headerless table whose first spec
        letter is ``a`` opens the NEXT sequential standard; one that resumes at a
        later letter continues the CURRENT standard.
      * sub-numbered items (``1.``, ``2.``) and wrapped rows accumulate as the
        current spec's body (never a new spec).
    """
    rows = list(table.rows)
    if not rows:
        return None

    header_std = None
    header_text = None
    for r in rows:
        for c in r.cells:
            m = _STD_HEADER_CELL_RE.match(_cell_text(c))
            if m:
                header_std = m.group("std")
                header_text = _cell_text(c)
                break
        if header_std:
            break

    if header_std is None:
        first_letter = _first_spec_letter(table)
        if first_letter is None:
            return None  # not a spec/standard table
        last = state.get("last_std")
        if first_letter == "a" and last:
            std = str(int(last) + 1)          # headerless → next sequential standard
        elif last:
            std = str(last)                   # headerless continuation of current std
        else:
            return None                        # no standard context yet → not a spec table
    else:
        std = header_std

    state["last_std"] = std
    state["seen_standard"] = True

    blocks: list[tuple] = []
    if header_std is not None and header_text:
        # Standard ROOT — the normative "shall" statement; routes to that
        # Standard's introduction (standardIntroductions[N]).
        blocks.append(("p", header_text, f"<p>{_esc(header_text)}</p>", False))

    cur_letter = None
    for r in rows:
        c0 = _cell_text(r.cells[0])
        if _STD_HEADER_CELL_RE.match(c0):
            continue  # header row already emitted
        # The content cell is the last non-marker cell (2-col: "a. | content";
        # 3-col: "a. | a. | content"). A merged single-column row keeps it in
        # cell 0 → no letter marker → continuation.
        content_cell = _content_cell(r)
        if len(r.cells) == 1:
            c0 = ""

        lm = _SPEC_LETTER_CELL_RE.match(c0)
        if lm:
            cur_letter = lm.group("letter").lower()
            prompt, body = _split_cell(content_cell)
            # Dotted std.letter form so _match_heading resolves BOTH std + spec.
            blocks.append(("p", f"{std}.{cur_letter}. {prompt}".rstrip(), "", False))
            blocks.extend(body)
            continue

        # Continuation row: a sub-numbered item ("1.", "2.") or a wrapped line —
        # accumulate as the CURRENT spec's body (never a new spec). Keep the
        # first-cell marker text (e.g. "2.") in front of the body line.
        if cur_letter is not None:
            paras = [p for p in content_cell.paragraphs if (p.text or "").strip()]
            for k, p in enumerate(paras):
                if _RESP_STANDALONE_RE.match(p.text or ""):
                    continue
                blk = _para_body_block(p)
                if k == 0 and c0:
                    blk = ("p", f"{c0} {blk[1]}".strip(), f"<p>{_esc(c0)} </p>{blk[2]}", blk[3])
                blocks.append(blk)

    return blocks


def _decompose_intro_table(table) -> list[tuple]:
    """Decompose a FRONT-matter table (Glossary, "General Program
    Characteristics" numbered prompts) into intro blocks. These sit before the
    first Standard, so the walker's introduction region routes them to the
    Document Introduction — never to a Standard."""
    rows = list(table.rows)
    if not rows:
        return []
    blocks: list[tuple] = []
    first_cell = _cell_text(rows[0].cells[0])
    if _GLOSSARY_RE.search(first_cell):
        # A. <Title> matches the walker's intro heading pattern, so the glossary
        # opens an intro section and its term rows land under it.
        blocks.append(("p", "A. Glossary of Terms", "<p><strong>Glossary of Terms</strong></p>", False))
        for r in rows:
            term = _cell_text(r.cells[0])
            defn = _cell_text(r.cells[1]) if len(r.cells) > 1 else ""
            if _GLOSSARY_RE.search(term) or (not term and not defn):
                continue
            line = f"{term}: {defn}".strip(": ").strip()
            if line:
                blocks.append(("p", line, f"<p>{_esc(line)}</p>", False))
        return blocks

    # General Program Characteristics: numbered/lettered prompts with Responses.
    for r in rows:
        c0 = _cell_text(r.cells[0])
        content_cell = r.cells[1] if len(r.cells) > 1 else r.cells[0]
        c1 = _cell_text(content_cell)
        if len(r.cells) == 1:
            c0 = ""
        if not c0 and not c1:
            continue
        if c0 and (_INTRO_NUM_CELL_RE.match(c0) or _SPEC_LETTER_CELL_RE.match(c0)):
            prompt, body = _split_cell(content_cell)
            blocks.append(("p", f"{c0} {prompt}".strip(), "", False))
            blocks.extend(body)
        else:
            for p in content_cell.paragraphs:
                if (p.text or "").strip():
                    blocks.append(_para_body_block(p))
    return blocks


def _run_inline_html(run_el, part) -> str:
    """Text (with bold/italic) + any inline images for one ``<w:r>`` run."""
    out: list[str] = []
    rpr = run_el.find(f"{_W}rPr")
    bold = rpr is not None and rpr.find(f"{_W}b") is not None
    ital = rpr is not None and rpr.find(f"{_W}i") is not None
    text = _esc("".join(t.text or "" for t in run_el.iter(f"{_W}t")))
    if text:
        if ital:
            text = f"<em>{text}</em>"
        if bold:
            text = f"<strong>{text}</strong>"
        out.append(text)
    # Inline images: <a:blip r:embed="rIdN"> → the related image part.
    for blip in run_el.iter(f"{_A}blip"):
        rid = blip.get(f"{_R}embed")
        if not rid:
            continue
        try:
            image_part = part.related_parts[rid]
            blob = image_part.blob
            if len(blob) > _MAX_INLINE_IMAGE_BYTES:
                out.append("<p><em>[image omitted — too large to inline; "
                           "attach the original via the Supporting File Library]</em></p>")
                continue
            ct = getattr(image_part, "content_type", "image/png") or "image/png"
            b64 = base64.b64encode(blob).decode("ascii")
            out.append(f'<img src="data:{ct};base64,{b64}" alt=""/>')
        except Exception:
            pass
    return "".join(out)


def _paragraph_to_html(p) -> str:
    """Inner HTML for one paragraph, preserving HYPERLINKS and inline IMAGES
    (the template's plain-text walk dropped both — the #1 fidelity complaint).
    """
    part = p.part
    out: list[str] = []
    for child in p._p:
        tag = child.tag
        if tag == f"{_W}hyperlink":
            rid = child.get(f"{_R}id")
            url = None
            if rid:
                try:
                    rel = part.rels[rid]
                    url = rel.target_ref if rel.is_external else None
                except Exception:
                    url = None
            inner = "".join(_run_inline_html(r, part) for r in child.findall(f"{_W}r"))
            out.append(f'<a href="{_esc(url)}">{inner}</a>' if url else inner)
        elif tag == f"{_W}r":
            out.append(_run_inline_html(child, part))
    return "".join(out)


def _iter_block_items(doc):
    """Yield ``("p", text, html)`` and ``("table", html, text)`` for every
    paragraph and table in the document, IN DOCUMENT ORDER. ``doc.paragraphs``
    alone drops tables AND flattens links/images to plain text; we walk the body
    XML so embedded tables land under the right section and the per-paragraph
    HTML keeps hyperlinks + inline images.
    """
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.table import Table as _Table
    from docx.text.paragraph import Paragraph as _Paragraph

    # Tabular official template? Only then do we decompose per-Standard response
    # tables + route front-matter tables to the Introduction. A paragraph
    # template (Kennesaw) or a doc with only evidence tables has no standard
    # tables, so this whole branch is skipped and tables stay opaque evidence —
    # existing behavior preserved.
    has_standard_tables = any(_looks_like_standard_table(t) for t in doc.tables)
    state = {"last_std": None, "seen_standard": False}

    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            para = _Paragraph(child, doc)
            # A paragraph "Standard N:" heading (paragraph template) also marks
            # that we've entered the standards region — so later evidence tables
            # aren't mistaken for front matter.
            if _STD_HEADER_CELL_RE.match((para.text or "").strip()):
                state["seen_standard"] = True
            # CR-062 — flag list-item paragraphs (Word numbering) so the body
            # walk can keep them as <li> inside a <ul> instead of flattening
            # every list bullet into its own <p> (Monica: "this should be in a
            # list form, but this is coming out in a paragraph form").
            yield ("p", para.text or "", _paragraph_to_html(para), _is_list_item(child))
        elif isinstance(child, CT_Tbl):
            t = _Table(child, doc)
            if has_standard_tables:
                std_blocks = _decompose_response_table(t, state)
                if std_blocks is not None:
                    yield from std_blocks
                    continue
                if not state["seen_standard"]:
                    # Front-matter table (Glossary / General Program
                    # Characteristics) → Document Introduction.
                    yield from _decompose_intro_table(t)
                    continue
            # Non-standard table (or paragraph template): opaque evidence table.
            yield ("table", _table_to_html(t), _table_to_text(t))


def _is_list_item(p_el) -> bool:
    """True when a paragraph carries Word list numbering (``w:pPr/w:numPr``)."""
    pPr = p_el.find(f"{_W}pPr")
    return pPr is not None and pPr.find(f"{_W}numPr") is not None


def _is_placeholder(body_paragraphs: list[str]) -> bool:
    """A section is a placeholder when every non-empty body paragraph
    matches the ``_PLACEHOLDER_RE`` pattern, or when the body is empty.

    ``See Appendix …`` lines alone are NOT placeholders — they're
    legitimate references — but a body of ONLY a "See Appendix" pointer
    AND nothing else is also treated as unwritten content for preview
    purposes (the institution hasn't actually authored a response yet).
    """
    real = [p.strip() for p in body_paragraphs if p.strip()]
    if not real:
        return True
    non_placeholder = [
        p for p in real
        if not _PLACEHOLDER_RE.match(p) and not p.lower().startswith("see appendix")
    ]
    return not non_placeholder


# ---------------------------------------------------------------- public API


def walk_template_paragraphs(
    paragraph_texts: list[str],
) -> list[TemplateSection]:
    """Pure-function walker over a list of paragraph strings (no tables).

    Thin wrapper over :func:`walk_template_blocks` — kept for the unit tests
    and any caller that only has paragraph text.
    """
    return walk_template_blocks([("p", t) for t in paragraph_texts])


def _is_real_break(
    hit: tuple[str | None, str | None],
    text: str,
    current_spec_section: "TemplateSection | None",
    in_introduction: bool,
    has_standards_region: bool,
) -> bool:
    """Decide whether a heading-looking line actually STARTS a new section.

    CR-061 — the governing import rule is *"write everything until the next
    break"*: a spec owns all content from its heading until the **next spec or
    standard**, not until the next numbered line. Inside the standards region,
    a spec's body routinely contains sub-numbered items ("3b 1.", "4a 3."),
    bare list numbers ("1.", "2." — advisory-board minutes, assessment lists),
    and lettered items that previously each opened a NEW section — fragmenting
    one spec across many sections that the matcher then assigned independently
    (or dropped at low confidence). The fix: only a forward spec/standard
    advance breaks; everything else accumulates as the current spec's body.
    """
    # The consolidation rule only applies inside the standards region of a full
    # template. The Introduction region keeps its own segmentation (its "1." /
    # "2a." prompts ARE its structure), and a bare fragment with no standards
    # region at all (e.g. unit-test snippets, partial pastes) keeps the legacy
    # behavior where every numbered heading cuts.
    if in_introduction or not has_standards_region:
        return True
    std_hint, spec_hint = hit
    # A Standard root ("Standard 4:" / bare "Standard 4") or a lettered group
    # header ("A. Institutional Requirements") sits BETWEEN standards — always a
    # break (the prior spec is already complete).
    if re.match(r"^\s*Standard\b", text, re.IGNORECASE):
        return True
    if re.match(r"^\s*[A-Z]\.\s+[A-Z]", text):
        return True
    # A bare std-only number ("1.", "2.") with no spec letter is a sub-item of
    # the current spec — accumulate, don't break.
    if spec_hint is None:
        return False
    # A spec heading ("3b 1." etc.) that names the SAME spec already being
    # processed is a sub-item — accumulate. Only a DIFFERENT spec breaks.
    if current_spec_section is not None:
        cur = (current_spec_section.standard_hint, current_spec_section.spec_hint)
        if cur == (std_hint, spec_hint):
            return False
    return True


def walk_template_blocks(blocks: list[tuple]) -> list[TemplateSection]:
    """Walk an in-document-order block stream into sections.

    Each block is ``("p", text)`` or ``("table", html, text)``. Paragraphs cut
    and fill sections exactly as before; a table attaches to the CURRENT
    section's ``evidence_tables`` (so embedded rosters/grids stay with their
    spec instead of being dropped). ``See Appendix …`` lines are also recorded
    on the section as import reminders.
    """
    sections: list[TemplateSection] = []
    current: TemplateSection | None = None
    # The spec/subspec CURRENTLY being processed — tables and in-prose evidence
    # are supporting evidence of THIS spec, even when they fall under a stray
    # list item ("2. The committee met…") that looks like a section root. Reset
    # only at a true Standard root (a new standard with no spec yet).
    current_spec_section: TemplateSection | None = None
    # Everything before the first Standard / Roman part header is the document
    # Introduction (ends with the glossary of terms). Only gate on this when the
    # document ACTUALLY has a standards region — otherwise a bare fragment of
    # "1." / "2a." headings (no Standard) keeps its spec hints (old behavior).
    text_blocks = [b[1] for b in blocks if b and b[0] == "p"]
    has_standards_region = any(
        _STANDARDS_REGION_RE.match((t or "").strip()) for t in text_blocks
    )
    in_introduction = has_standards_region

    for idx, block in enumerate(blocks):
        if block and block[0] == "table":
            # The table is evidence of the spec currently being processed (carry
            # forward across stray list-item "headings"); fall back to the
            # current section (e.g. an intro table with no spec yet).
            target = current_spec_section if current_spec_section is not None else current
            if target is not None:
                target.evidence_tables.append({"html": block[1], "text": block[2]})
            # Keep the table INLINE in the narrative too (fidelity — "the table
            # is missing"). It lands in the section it physically sits under.
            if current is not None and block[1]:
                current.body_html_parts.append(block[1])
            continue

        text = (block[1] if block else "").strip()
        para_html = block[2] if (block and len(block) > 2) else ""
        if not text:
            # An image-only / blank paragraph has no text but may carry an
            # embedded image — keep it in the narrative instead of dropping it.
            if current is not None and "<img" in para_html:
                current.body_html_parts.append(f"<p>{para_html}</p>")
            continue

        # Skip front-matter lines (template title, "Introduction and
        # Instructions:", "Table of Contents"). They are not section
        # cuts AND they should not be accumulated into the body of the
        # previous section. Dropping them entirely is cleanest.
        if _is_front_matter(text):
            continue

        # Leaving the Introduction: the first Standard (or Roman part header)
        # ends it. Detected BEFORE heading classification so the Standard line
        # itself is NOT tagged as introduction.
        if in_introduction and _STANDARDS_REGION_RE.match(text):
            in_introduction = False

        hit = _match_heading(text)
        if hit is not None and _is_real_break(
            hit, text, current_spec_section, in_introduction, has_standards_region
        ):
            # Close the previous section.
            if current is not None:
                current.placeholder = _is_placeholder(current.body_paragraphs)
                sections.append(current)
            std_hint, spec_hint = hit
            if in_introduction:
                # Intro numbering ("1.", "2a.", "B. Glossary") is the intro's
                # own, NOT Standard specs — drop the hints so nothing buckets
                # under a Standard; the section routes to introduction:document.
                std_hint, spec_hint = None, None
            current = TemplateSection(
                paragraph_index=idx,
                heading=text,
                body_paragraphs=[],
                standard_hint=std_hint,
                spec_hint=spec_hint,
                placeholder=False,  # set on close
                is_introduction=in_introduction,
            )
            # Track the spec currently being processed for evidence attribution.
            is_std_root = (
                not in_introduction
                and std_hint is not None
                and spec_hint is None
                and bool(re.match(r"^\s*Standard\b", text, re.IGNORECASE))
            )
            if is_std_root:
                current_spec_section = None  # new standard, no spec yet
            elif not in_introduction and std_hint and spec_hint:
                current_spec_section = current  # this spec is now active
            continue

        if current is None:
            # Paragraphs before the first heading are preamble — drop them.
            # The template's actual content always starts under a heading.
            continue

        # Record any "See Appendix …" references as import reminders — under the
        # spec currently being processed (kept in the body prose too).
        refs = _extract_appendix_refs(text)
        if refs:
            (current_spec_section or current).appendix_refs.extend(refs)

        # Inside a section. Strip ``Response:`` markers; if the line is
        # only the marker, skip it entirely (the body comes on the next
        # paragraph). Otherwise keep what remains.
        stripped = _strip_response_marker(text)
        if not stripped:
            continue
        current.body_paragraphs.append(stripped)
        # Rich-HTML narrative: prefer the paragraph's HTML (keeps hyperlinks /
        # inline images / bold-italic), with a leading "Response:" stripped to
        # match body_paragraphs; fall back to escaped plain text.
        para_html = block[2] if len(block) > 2 else ""
        is_list = block[3] if len(block) > 3 else False
        if para_html.strip():
            para_html = re.sub(r"^\s*(?:<[^>]+>\s*)*Response\s*:\s*", "", para_html, flags=re.IGNORECASE)
            current.body_html_parts.append(f"<li>{para_html}</li>" if is_list else f"<p>{para_html}</p>")
        else:
            current.body_html_parts.append(f"<li>{_esc(stripped)}</li>" if is_list else f"<p>{_esc(stripped)}</p>")

    if current is not None:
        current.placeholder = _is_placeholder(current.body_paragraphs)
        sections.append(current)

    return sections


def walk_template_docx(
    docx_path: str,
    base_id: str = "template",
) -> tuple[list[Section], list[TemplateSection]]:
    """Open a DOCX and walk it. Returns ``(sections, raw_template_sections)``.

    ``sections`` is the ``Section`` list the matcher consumes.
    ``raw_template_sections`` is the intermediate form, kept so the
    preview renderer can show placeholder/empty sections that the
    matcher never sees.
    """
    doc = Document(docx_path)
    blocks = list(_iter_block_items(doc))
    raw_sections = walk_template_blocks(blocks)

    sections: list[Section] = []
    for raw in raw_sections:
        # A Standard ROOT ("Standard 9: …the program shall…") carries its
        # normative statement IN THE HEADING, so its body is often empty — but
        # it still anchors that Standard's introduction. Emit it even when the
        # body is a placeholder; skip only genuinely-unwritten spec responses.
        is_std_root = bool(
            raw.standard_hint
            and not raw.spec_hint
            and not raw.is_introduction
            and re.match(r"^\s*Standard\b", raw.heading, re.IGNORECASE)
        )
        if raw.placeholder and not is_std_root:
            # Don't send unwritten / "Not applicable" sections through the
            # matcher — there's nothing to classify and we'd just waste
            # API calls (and add noise to the preview).
            continue
        md = f"# {raw.heading}\n\n{raw.body_text}"
        flags = _heuristic_flags(md)
        # Carry the heading-derived hints in flags so downstream coverage /
        # bucket-allocation can also surface them, not just the matcher.
        # Introduction-region sections carry NO standard/spec hint and a
        # document-introduction routing hint instead — the template pipeline
        # routes them to the documentIntroduction bucket rather than the matcher.
        if raw.is_introduction:
            flags["templateStandardHint"] = ""
            flags["templateSpecHint"] = ""
            flags["templateIntroductionHint"] = "introduction:document"
        elif raw.standard_hint and not raw.spec_hint and re.match(
            r"^\s*Standard\b", raw.heading, re.IGNORECASE
        ):
            # A Standard ROOT ("Standard 1: …the program shall…") — the
            # normative sentence + Context framing is NOT a spec response;
            # route it to that Standard's introduction (standardIntroductions[N]).
            flags["templateStandardHint"] = raw.standard_hint
            flags["templateSpecHint"] = ""
            flags["templateIntroductionHint"] = f"introduction:standard-{raw.standard_hint}"
        else:
            flags["templateStandardHint"] = raw.standard_hint or ""
            flags["templateSpecHint"] = raw.spec_hint or ""
        flags["templateHeading"] = raw.heading[:300]
        # The faithful narrative content: paragraphs with hyperlinks + inline
        # images and tables, in document order. The renderer/editor prefers
        # html_snippet when present (markdown stays plain for the matcher).
        body_html = raw.body_html
        sections.append(
            Section(
                id=f"{base_id}:tmpl:{uuid.uuid4().hex[:8]}",
                # Spec prompts / standard statements run long ("2b. Describe the
                # institutional context …") — a 200-char cap chopped them mid-word
                # ("text cut off"). 600 keeps the full prompt.
                heading=raw.heading[:600],
                heading_level=2,
                markdown=md,
                byte_offset_start=0,
                byte_offset_end=len(md),
                word_count=len(md.split()),
                contains_table=flags.get("containsTable", False),
                contains_image=("<img" in body_html) or flags.get("containsImage", False),
                has_resume_signals=flags.get("hasResumeSignals", False),
                has_syllabus_signals=flags.get("hasSyllabusSignals", False),
                splitter_tier="template",
                flags=flags,
                html_snippet=(body_html or None),
            )
        )
    return sections, raw_sections
