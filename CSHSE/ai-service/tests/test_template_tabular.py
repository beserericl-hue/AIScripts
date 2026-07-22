"""Tabular official CSHSE template — one table per Standard, with the
``Standard N:`` heading, lettered specs (a./b./…) and each ``Response:`` all
inside table CELLS. Several institutions (e.g. Anne Arundel CC) submit the
official template with its tables intact; Kennesaw removed them. This locks in:

  * the detector classifies a purely-tabular template as ``template`` (its
    signals live in cells, not paragraphs);
  * per-Standard response tables are decomposed into standard/spec sections;
  * a HEADERLESS continuation table opens the next sequential Standard;
  * a 3-column table (``a. | a. | <content>``) reads the content column;
  * all three answer shapes parse — ``Response:``, bare ``Response`` and
    NO marker (answer written directly under the prompt);
  * front-matter tables (Glossary, General Program Characteristics) route to the
    Introduction, never to a Standard.
"""
import os
import tempfile

import pytest
from docx import Document

from app.splitter.format_detector import detect_format
from app.splitter.template_walker import walk_template_docx


def _set_cell(cell, lines):
    """Write LINES as separate paragraphs in a cell (as a real docx does)."""
    cell.paragraphs[0].text = lines[0]
    for ln in lines[1:]:
        cell.add_paragraph(ln)


def _std_table(doc, header, rows, cols=2):
    """Add a per-Standard response table. header=None → headerless continuation.
    Each row is (letter_or_'', [content paragraphs]). Middle columns duplicate
    the letter marker (the 3-column ``a. | a. | content`` shape)."""
    body = ([("__hdr__", [header])] if header else []) + rows
    t = doc.add_table(rows=len(body), cols=cols)
    for i, (marker, lines) in enumerate(body):
        cells = t.rows[i].cells
        if marker == "__hdr__":
            for c in cells:
                c.text = lines[0]
            continue
        cells[0].text = marker
        for c in cells[1:-1]:
            c.text = marker
        _set_cell(cells[-1], lines)
    return t


@pytest.fixture
def tabular_docx():
    d = Document()
    # Front matter → Introduction
    g = d.add_table(rows=2, cols=2)
    g.rows[0].cells[0].text = "Glossary of Terms"
    g.rows[1].cells[0].text = "AACC"; g.rows[1].cells[1].text = "Anne Arundel Community College"
    gp = d.add_table(rows=2, cols=2)
    gp.rows[0].cells[0].text = "1."
    _set_cell(gp.rows[0].cells[1], ["Specify the degree.", "Response:", "Associate of Applied Sciences."])
    gp.rows[1].cells[0].text = "2."
    _set_cell(gp.rows[1].cells[1], ["Describe the institution.", "Response:", "AACC is a public two-year college."])

    # Standard 4 — the three answer shapes: 'Response:', bare 'Response', NO marker
    _std_table(d, "Standard 4: The program shall conduct evaluations.", [
        ("a.", ["Measurable student learning outcomes.", "Response:", "The program has clear SLOs and an assessment plan."]),
        ("b.", ["Formal program evaluation.", "Response", "A five-year evaluation cycle is in place."]),
        ("c.", ["Report to the public.", "", "The program publishes achievement data on its website."]),
    ])
    # Standard 5 — HEADERLESS continuation (opens the next sequential standard)
    _std_table(d, None, [
        ("a.", ["Admission policies.", "Response:", "Open-access admissions per college policy."]),
        ("b.", ["Referral policies.", "Response:", "Faculty use the Student Referral System."]),
    ])
    # Standard 6 — 3-column table (letter duplicated in cols 0 and 1)
    _std_table(d, "Standard 6: Faculty competencies.", [
        ("a.", ["Faculty credentials.", "Response:", "All faculty hold at least a master's degree."]),
    ], cols=3)

    f = tempfile.mktemp(suffix=".docx")
    d.save(f)
    yield f
    os.remove(f)


def test_tabular_template_is_detected_as_template(tabular_docx):
    v = detect_format(tabular_docx)
    assert v.format == "template", v.reasoning
    assert v.signals["response_marker_count"] >= 3


def test_tabular_specs_route_to_standards_and_intro(tabular_docx):
    sections, _ = walk_template_docx(tabular_docx)
    specs = {
        (s.flags.get("templateStandardHint"), s.flags.get("templateSpecHint")): s
        for s in sections
        if s.flags.get("templateSpecHint")
    }
    # Standard 4 all three answer shapes captured under the RIGHT spec.
    assert ("4", "a") in specs and ("4", "b") in specs and ("4", "c") in specs
    assert "clear SLOs" in specs[("4", "a")].markdown
    assert "five-year evaluation" in specs[("4", "b")].markdown          # bare 'Response'
    assert "publishes achievement data" in specs[("4", "c")].markdown    # NO marker
    # Standard 5 recovered from the HEADERLESS continuation table.
    assert ("5", "a") in specs and ("5", "b") in specs
    assert "Open-access" in specs[("5", "a")].markdown
    # Standard 6 read the CONTENT column of the 3-wide table.
    assert ("6", "a") in specs
    assert "master's degree" in specs[("6", "a")].markdown

    # Front matter → Document Introduction, NOT a Standard.
    intro = [s for s in sections if s.flags.get("templateIntroductionHint") == "introduction:document"]
    intro_text = "\n".join(s.markdown for s in intro)
    assert "Associate of Applied Sciences" in intro_text
    assert "Glossary" in intro_text or "Anne Arundel" in intro_text
    # The intro's "Specify the degree" prompt must NOT have leaked into a Standard spec.
    for (std, spec), s in specs.items():
        assert "Specify the degree" not in s.markdown
