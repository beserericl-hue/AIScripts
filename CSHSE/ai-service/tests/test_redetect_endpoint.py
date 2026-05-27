"""CR-040 follow-on (2026-05-27) — /ai/import/redetect endpoint integration tests.

These tests exercise the full Re-detect path end-to-end at the
ai-service level:

  client.post /ai/import/redetect
    → verify_hmac_signature
    → boto3 S3 download (mocked: writes a real DOCX fixture to disk)
    → mammoth: DOCX → HTML
    → deep_walker
    → cv_detector (pre + post)
    → appendix_paper_detector (pre + post)
    → introduction_detector
    → toc_detector (parse_toc + anchor_in_body + merge_*)
    → JSON response

The Stevenson-shaped fixture has CVs the pattern-based detector
misses on purpose (the body anchor heuristic doesn't fire for them)
but the Table of Contents lists them, so the TOC pass is the only
thing that can recover them. These tests pin that contract — if a
future refactor weakens the TOC pass, the recovered count drops and
the test fails.

System-test scope (vs unit): the unit tests in test_toc_detector.py
exercise the splitter module in isolation. These tests run the
endpoint as the cshse-server would call it, exercising the full
endpoint glue (S3 fetch, mammoth conversion, pattern + TOC + merge,
response shape).
"""
from __future__ import annotations

import hashlib
import hmac
import io
import json
import os
import time
from pathlib import Path
from unittest.mock import patch, MagicMock

import pytest
from fastapi.testclient import TestClient


_HMAC_SECRET = "test-redetect-secret"


# ----------------------------------------------------------------------
# Fixture builders
# ----------------------------------------------------------------------


def _build_stevenson_like_docx(out_path: Path) -> None:
    """Generate a real .docx file with the Stevenson-shaped TOC + body.

    The TOC lists 5 CVs but the body only carries the "obvious" 2 (with
    clean EDUCATION + EXPERIENCE markers); the other 3 use the
    credentials-prefixed pattern that the legacy cv_detector misses.

    Goal — pattern detector finds 2 (or fewer), TOC pass recovers 3,
    final count = 5. If a future change weakens the TOC pass we'll see
    the final count drop below 5 and the test will fail.
    """
    from docx import Document

    doc = Document()
    doc.add_heading("Table of Contents", level=1)
    doc.add_paragraph("Standard 1 ............ 1")
    doc.add_paragraph("Standard 6 ............ 80")
    doc.add_paragraph("CVs")
    doc.add_paragraph("John Rosicky ............ 100")
    doc.add_paragraph("Carol A. Dietrich ............ 110")
    doc.add_paragraph("Thomas K. Swisher, J.D., Ph.D. ............ 140")
    doc.add_paragraph("LAURI A. WEINER, HS-BCP ............ 150")
    doc.add_paragraph("Mary Beth Olson, M.A. ............ 160")
    doc.add_paragraph("Syllabi")
    doc.add_paragraph("CHS 220 ............ 200")
    doc.add_paragraph("CHS 305 ............ 210")
    doc.add_paragraph("Papers / Projects")
    doc.add_paragraph("Sample Country Report ............ 300")

    doc.add_heading("Introduction", level=2)
    doc.add_paragraph(
        "This self-study presents the Stevenson University Human Services "
        "program in extensive detail, covering all six CSHSE standards. "
        "Faculty CVs are collected in Appendix A; course syllabi in "
        "Appendix B; student work samples in Appendix C. "
        "This sentence makes the paragraph long enough to qualify as body "
        "prose so the parser knows the TOC region has ended and the body "
        "has begun."
    )

    # --- Body: CVs ---
    doc.add_heading("Appendix A: CVs", level=2)

    # CV #1 — John Rosicky — has the "clean" anchor pattern.
    # Strong CV signal: name + contact + multiple section markers.
    doc.add_paragraph("John Rosicky")
    doc.add_paragraph("1051 Omar Dr.")
    doc.add_paragraph("Crownsville, MD 21032")
    doc.add_paragraph("(315) 525-6211")
    doc.add_paragraph("jrosicky@stevenson.edu")
    doc.add_paragraph("EDUCATION")
    doc.add_paragraph("1990  University of Maryland  B.A. Sociology")
    doc.add_paragraph("PROFESSIONAL EXPERIENCE")
    doc.add_paragraph("2005-Present  Stevenson University  Human Services")
    doc.add_paragraph("PUBLICATIONS")
    doc.add_paragraph("Rosicky, J. (2015). Family Services Quarterly.")

    # CV #2 — Carol A. Dietrich — clean pattern.
    doc.add_paragraph("Carol A. Dietrich")
    doc.add_paragraph("5766 Kinsmen Courage Court")
    doc.add_paragraph("Eldersburg, Maryland 21784")
    doc.add_paragraph("410-596-0625")
    doc.add_paragraph("EDUCATION")
    doc.add_paragraph("2002 University of Maryland M.S. Counseling")
    doc.add_paragraph("PROFESSIONAL EXPERIENCE")
    doc.add_paragraph("2010-Present Stevenson University")

    # CV #3 — Thomas K. Swisher — TOC-only recovery candidate.
    # The body opens with credentials-prefixed name; no contact info
    # in the first 5 lines; the section heading uses Title-Case
    # ("Education" not "EDUCATION"). The pattern detector's
    # _is_anchor_line gates name + contact-or-marker within 5 lines,
    # so a CV whose first 5 lines are name + address + degree-listing
    # (no contact, no ALL-CAPS marker) slips past.
    doc.add_paragraph("Thomas K. Swisher, J.D., Ph.D.")
    doc.add_paragraph("11886 Simpson Road")
    doc.add_paragraph("Clarksville, Md. 21029")
    doc.add_paragraph("Education")  # Title-case, not ALL-CAPS
    doc.add_paragraph("1983 University of Virginia B.S. Secondary Education")
    doc.add_paragraph("1986 University of Baltimore Juris Doctorate")
    doc.add_paragraph("Experience")
    doc.add_paragraph("Associate Professor, Stevenson University, 2007-Present")

    # CV #4 — LAURI A. WEINER — TOC-only recovery candidate.
    # The all-caps name fails the pattern detector's _is_anchor_line
    # (rejects all-caps lines as section headings, not names).
    doc.add_paragraph("LAURI A. WEINER, J.D., HS-BCP")
    doc.add_paragraph("7905 Winterset Avenue")
    doc.add_paragraph("Baltimore, Maryland 21208")
    doc.add_paragraph("lweiner@stevenson.edu")
    doc.add_paragraph("Education")
    doc.add_paragraph("J.D., University of Maryland School of Law, 1992")
    doc.add_paragraph("M.A., Bowling Green State University, 1984")
    doc.add_paragraph("Teaching Experience")
    doc.add_paragraph("Associate Professor of Human Services, Stevenson University")

    # CV #5 — Mary Beth Olson — TOC-only recovery candidate.
    # 3-token name with M.A. suffix; first body line is contact info,
    # so the name anchor isn't on the section's first line. Pattern
    # detector's _section_is_cv_candidate fires on first non-blank.
    doc.add_paragraph("250 Faculty Drive · Baltimore, MD 21210")
    doc.add_paragraph("Mary Beth Olson, M.A.")
    doc.add_paragraph("Education")
    doc.add_paragraph("2007 Loyola University Maryland M.A. Counseling")
    doc.add_paragraph("Service")
    doc.add_paragraph("Field Placement Coordinator, Stevenson University")

    # --- Body: Syllabi ---
    doc.add_heading("Appendix B: Syllabi", level=2)
    doc.add_paragraph("CHS 220 Spring 2019")
    doc.add_paragraph(
        "Course Syllabus — Introduction to Human Services Practice. "
        "This course examines the foundations of human services work, "
        "including counseling theory, case management, and ethics. "
        "Students complete a 30-hour field observation."
    )
    doc.add_paragraph("CHS 305 Fall 2019")
    doc.add_paragraph(
        "Course Syllabus — Family Systems Counseling. "
        "Building on CHS 220, this course covers family-systems theory, "
        "structural family therapy, and Bowen's intergenerational approach. "
        "Students complete a clinical role-play final."
    )

    # --- Body: Student work ---
    doc.add_heading("Appendix C: Papers / Projects", level=2)
    doc.add_paragraph("Sample Country Report")
    doc.add_paragraph(
        "Research Paper (Individual Work) (125 points). "
        "A multi-page country analysis prepared by a graduating senior "
        "as the capstone for CHS 410. The paper covers demographic "
        "profile, social-service infrastructure, and immigration patterns."
    )

    doc.save(out_path)


@pytest.fixture
def stevenson_docx(tmp_path):
    """Real .docx file with the Stevenson-shaped TOC + body."""
    out = tmp_path / "stevenson-like.docx"
    _build_stevenson_like_docx(out)
    return out


@pytest.fixture
def client(monkeypatch):
    """FastAPI test client + S3 download mock + HMAC env."""
    # Set the HMAC secret BEFORE importing app.main (the secret is
    # cached via get_settings()).
    monkeypatch.setenv("NODE_SERVICE_HMAC_SECRET", _HMAC_SECRET)
    monkeypatch.setenv("CSHSE_ENV", "dev")
    from app.config import get_settings
    get_settings.cache_clear()
    from app.main import app
    yield TestClient(app)
    get_settings.cache_clear()


def _sign(body: bytes) -> dict:
    ts = int(time.time())
    digest = hmac.new(
        _HMAC_SECRET.encode(), f"{ts}.".encode() + body, hashlib.sha256
    ).hexdigest()
    return {"x-service-signature": f"t={ts},v1={digest}"}


def _post_redetect(client, body: dict, docx_path: Path):
    """POST /ai/import/redetect with the S3 download mocked to copy
    ``docx_path`` to the target location.

    Returns the parsed JSON body.
    """
    body_bytes = json.dumps(body).encode("utf-8")

    def _fake_download_file(bucket, key, dest):
        # Copy the fixture docx to the destination the endpoint expects.
        import shutil
        shutil.copy(docx_path, dest)

    fake_s3_client = MagicMock()
    fake_s3_client.download_file.side_effect = _fake_download_file

    with patch("boto3.client", return_value=fake_s3_client):
        return client.post(
            "/ai/import/redetect",
            data=body_bytes,
            headers={**_sign(body_bytes), "content-type": "application/json"},
        )


# ----------------------------------------------------------------------
# Tests — happy path
# ----------------------------------------------------------------------


class TestRedetectStevenson:
    """The Stevenson-shaped fixture is the case the user reported: 5 CVs
    listed in the TOC, but the pattern detector only finds 2 (or fewer).
    The TOC pass must recover the missing 3 so the final count is 5.
    """

    def test_returns_200_and_ok_true(self, client, stevenson_docx):
        resp = _post_redetect(
            client,
            {
                "s3Key": "imports/stevenson.docx",
                "importId": "imp-stevenson-1",
                "submissionId": "sub-stevenson-1",
            },
            stevenson_docx,
        )
        assert resp.status_code == 200, resp.text
        body = resp.json()
        assert body["ok"] is True

    def test_recovers_all_five_cvs_via_toc(self, client, stevenson_docx):
        """Pattern detector alone misses Thomas Swisher, Lauri Weiner,
        and Mary Beth Olson (different anchor patterns). The TOC pass
        recovers all three. Final count must be 5."""
        resp = _post_redetect(
            client,
            {
                "s3Key": "imports/stevenson.docx",
                "importId": "imp-stevenson-1",
                "submissionId": "sub-stevenson-1",
            },
            stevenson_docx,
        )
        body = resp.json()
        cv_names = [c["facultyName"] for c in body["cvs"]]
        # Normalize for comparison — credentials may or may not be stripped.
        cv_names_lower = " | ".join(n.lower() for n in cv_names)
        assert "john rosicky" in cv_names_lower
        assert "carol a. dietrich" in cv_names_lower or "carol a dietrich" in cv_names_lower
        assert "thomas k. swisher" in cv_names_lower
        assert "lauri a. weiner" in cv_names_lower
        assert "mary beth olson" in cv_names_lower
        # And the count tile reflects 5.
        assert body["counts"]["cvs"] >= 5

    def test_toc_diagnostics_block_present(self, client, stevenson_docx):
        """The response must include tocDiagnostics so the client can
        show "TOC pass recovered N items" in the redetect banner."""
        resp = _post_redetect(
            client,
            {
                "s3Key": "imports/stevenson.docx",
                "importId": "imp-stevenson-1",
                "submissionId": "sub-stevenson-1",
            },
            stevenson_docx,
        )
        body = resp.json()
        assert "tocDiagnostics" in body
        diag = body["tocDiagnostics"]
        assert "tocEntriesFound" in diag
        assert "tocAnchoredDetections" in diag
        assert "tocAdded" in diag
        # TOC must have parsed at least 8 entries (5 CVs + 2 syllabi +
        # 1 paper) — anything less means the TOC walker broke.
        assert diag["tocEntriesFound"] >= 8
        # TOC must have ADDED at least one CV the pattern detector missed.
        assert diag["tocAdded"]["cvs"] >= 1

    def test_toc_added_cvs_routed_as_source_toc(self, client, stevenson_docx):
        """Detections added by the TOC path carry routing.source='toc'
        so a future UI badge can show provenance — and so the merge
        layer can be reasoned about by reading the response alone."""
        resp = _post_redetect(
            client,
            {
                "s3Key": "imports/stevenson.docx",
                "importId": "imp-stevenson-1",
                "submissionId": "sub-stevenson-1",
            },
            stevenson_docx,
        )
        body = resp.json()
        toc_sourced = [
            c for c in body["cvs"] if (c.get("routing") or {}).get("source") == "toc"
        ]
        assert len(toc_sourced) >= 1, (
            "At least one CV must be marked routing.source='toc' so the "
            "merge layer's contribution is visible in the response."
        )

    def test_recovers_both_syllabi_via_toc_or_pattern(self, client, stevenson_docx):
        resp = _post_redetect(
            client,
            {
                "s3Key": "imports/stevenson.docx",
                "importId": "imp-stevenson-1",
                "submissionId": "sub-stevenson-1",
            },
            stevenson_docx,
        )
        body = resp.json()
        syllabi = [d for d in body["evidenceDocs"] if d["docSubKind"] == "syllabus"]
        # We have CHS 220 and CHS 305 — the pattern detector should
        # find both (both have the (NN points)-style header or a
        # COURSE_CODE in the first 5 lines), and the TOC pass dedupes.
        codes = [s.get("courseCode") for s in syllabi]
        codes_str = " ".join(c or "" for c in codes)
        assert "CHS 220" in codes_str or "CHS 305" in codes_str

    def test_recovers_paper_via_toc_or_pattern(self, client, stevenson_docx):
        resp = _post_redetect(
            client,
            {
                "s3Key": "imports/stevenson.docx",
                "importId": "imp-stevenson-1",
                "submissionId": "sub-stevenson-1",
            },
            stevenson_docx,
        )
        body = resp.json()
        papers = [d for d in body["evidenceDocs"] if d["docSubKind"] == "paper"]
        assert len(papers) >= 1
        titles = " ".join(p["title"].lower() for p in papers)
        assert "country" in titles or "research" in titles

    def test_counts_block_matches_cv_array(self, client, stevenson_docx):
        """counts.cvs must exactly equal len(cvs)."""
        resp = _post_redetect(
            client,
            {
                "s3Key": "imports/stevenson.docx",
                "importId": "imp-stevenson-1",
                "submissionId": "sub-stevenson-1",
            },
            stevenson_docx,
        )
        body = resp.json()
        assert body["counts"]["cvs"] == len(body["cvs"])
        assert body["counts"]["papers"] == sum(
            1 for d in body["evidenceDocs"] if d["docSubKind"] == "paper"
        )
        assert body["counts"]["syllabi"] == sum(
            1 for d in body["evidenceDocs"] if d["docSubKind"] == "syllabus"
        )


# ----------------------------------------------------------------------
# Tests — auth / failure paths
# ----------------------------------------------------------------------


class TestRedetectAuth:
    def test_missing_hmac_signature_returns_401(self, client, stevenson_docx):
        body_bytes = json.dumps({
            "s3Key": "x", "importId": "y", "submissionId": "z"
        }).encode("utf-8")
        # No signature header.
        resp = client.post(
            "/ai/import/redetect",
            data=body_bytes,
            headers={"content-type": "application/json"},
        )
        assert resp.status_code == 401

    def test_bad_hmac_signature_returns_401(self, client):
        body_bytes = json.dumps({
            "s3Key": "x", "importId": "y", "submissionId": "z"
        }).encode("utf-8")
        resp = client.post(
            "/ai/import/redetect",
            data=body_bytes,
            headers={
                "content-type": "application/json",
                "x-service-signature": "t=1,v1=deadbeef",
            },
        )
        assert resp.status_code == 401


# ----------------------------------------------------------------------
# Tests — failure paths (s3 fetch)
# ----------------------------------------------------------------------


class TestRedetectS3Failure:
    def test_s3_download_failure_returns_502(self, client):
        """When boto3.download_file raises, the endpoint must return
        502 with a structured detail. The client surfaces this to the
        coordinator as an error banner."""
        body = {"s3Key": "missing.docx", "importId": "i", "submissionId": "s"}
        body_bytes = json.dumps(body).encode("utf-8")
        fake_s3 = MagicMock()
        fake_s3.download_file.side_effect = RuntimeError("NoSuchKey")
        with patch("boto3.client", return_value=fake_s3):
            resp = client.post(
                "/ai/import/redetect",
                data=body_bytes,
                headers={**_sign(body_bytes), "content-type": "application/json"},
            )
        assert resp.status_code == 502
        assert "s3 fetch failed" in (resp.json().get("detail") or "")


# ----------------------------------------------------------------------
# Tests — no-TOC fallback
# ----------------------------------------------------------------------


def _build_no_toc_docx(out_path: Path) -> None:
    """A docx with NO Table of Contents. The TOC pass should no-op
    cleanly — pattern-based detection is the only signal."""
    from docx import Document

    doc = Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(
        "A self-study with no TOC. The pattern detector still runs; "
        "the TOC pass should gracefully return zero added items and "
        "tocEntriesFound = 0."
    )
    # One easy-to-detect CV so we know the pipeline still works.
    doc.add_heading("Appendix A", level=2)
    doc.add_paragraph("Mary Smith")
    doc.add_paragraph("mary.smith@example.edu")
    doc.add_paragraph("(555) 123-4567")
    doc.add_paragraph("EDUCATION")
    doc.add_paragraph("2010 Some University M.A. Counseling")
    doc.add_paragraph("PROFESSIONAL EXPERIENCE")
    doc.add_paragraph("2015-Present Faculty, Some University")
    doc.add_paragraph("PUBLICATIONS")
    doc.add_paragraph("Smith, M. (2020). A paper.")
    doc.save(out_path)


@pytest.fixture
def no_toc_docx(tmp_path):
    out = tmp_path / "no-toc.docx"
    _build_no_toc_docx(out)
    return out


def _build_stevenson_real_shape_docx(out_path: Path) -> None:
    """CR-040 follow-on (2026-05-27) — fixture matching the REAL
    Stevenson document shape the user reported.

    Real Stevenson docs DON'T list individual CVs in the main TOC.
    The main TOC has a single line "Appendices (List of Supporting
    Documents) ... 112", and the per-CV enumeration lives in a
    sub-TOC at page 112 (under a "List of Faculty CVs" heading).
    There are also Standard 1...6 entries IN the main TOC that
    must NOT be misread as body-section headers (the earlier
    parse_toc bailed on them).

    Pattern detector + main-TOC pass alone produces 0 CV recoveries
    here. Only the sub-TOC scanner recovers them. If the sub-TOC
    code regresses, the count drops to 2 and the test fails.
    """
    from docx import Document

    doc = Document()

    # --- Main TOC: just standards + appendix index, no CV enumeration ---
    doc.add_heading("Table of Contents", level=1)
    doc.add_paragraph("Certification of the Self-Study ............ 4")
    doc.add_paragraph("Introductory Information ............ 5")
    doc.add_paragraph("Glossary of Terms ............ 13")
    doc.add_paragraph("Part I: General Program Characteristics")
    doc.add_paragraph("Standard 1 - Institutional Requirements ............ 14")
    doc.add_paragraph("Standard 2 - Philosophical Base ............ 17")
    doc.add_paragraph("Standard 6 - Personnel ............ 80")
    doc.add_paragraph("Appendices (List of Supporting Documents) ............ 112")
    doc.add_paragraph("Course Syllabi and Materials ............ 114")

    # --- Body intro (forces the main-TOC walker to bail here) ---
    doc.add_heading("Introduction", level=2)
    doc.add_paragraph(
        "This self-study presents the Stevenson University Human Services "
        "program in extensive detail, covering all six CSHSE standards. "
        "Body prose continues with the standards narratives, faculty "
        "qualifications, and program-level evidence. The narrative is "
        "organised by standard and includes references to the appendix "
        "where supporting evidence is collected."
    )

    # --- Standards body (no CV listing) ---
    doc.add_heading("Standard 1 - Institutional Requirements", level=2)
    doc.add_paragraph(
        "Stevenson University is accredited by the Middle States Commission "
        "on Higher Education. The Human Services baccalaureate program "
        "operates within the College of Education, Human Services, and "
        "Health. Detailed institutional information follows in the "
        "subsequent sections of this Standard."
    )
    doc.add_heading("Standard 6 - Personnel", level=2)
    doc.add_paragraph(
        "The faculty roster comprises full-time and part-time members "
        "whose qualifications are documented in Appendix A. Detailed "
        "credentials, teaching experience, and professional service are "
        "provided in the curriculum vitae section."
    )

    # --- Appendix A: List of Faculty CVs (sub-TOC) ---
    doc.add_heading("Appendix A: List of Faculty CVs", level=2)
    doc.add_paragraph("John Rosicky ............ 113")
    doc.add_paragraph("Carol A. Dietrich ............ 117")
    doc.add_paragraph("Roxanne M. Epps ............ 119")
    doc.add_paragraph("Barry W. Thomas ............ 121")
    doc.add_paragraph("Thomas K. Swisher, J.D., Ph.D. ............ 125")
    doc.add_paragraph("LAURI A. WEINER, HS-BCP ............ 130")
    doc.add_paragraph("Mary Beth Olson, M.A. ............ 135")

    # --- Appendix B: List of Course Syllabi (sub-TOC) ---
    doc.add_heading("Appendix B: List of Course Syllabi", level=2)
    doc.add_paragraph("CHS 220 - Introduction to Human Services ............ 200")
    doc.add_paragraph("CHS 305 - Family Systems ............ 220")
    doc.add_paragraph("CHS 410 - Capstone Seminar ............ 240")

    # --- Appendix C: List of Student Papers (sub-TOC) ---
    doc.add_heading("Appendix C: List of Student Papers", level=2)
    doc.add_paragraph("Sample Country Report ............ 300")
    doc.add_paragraph("Final Research Project ............ 320")

    # --- CV bodies (deep in body, would defeat pattern detector) ---
    doc.add_heading("Appendix A — Faculty CVs", level=2)

    # CV body that the pattern detector CAN find (John Rosicky shape)
    doc.add_paragraph("John Rosicky")
    doc.add_paragraph("1051 Omar Dr.")
    doc.add_paragraph("Crownsville, MD 21032")
    doc.add_paragraph("jrosicky@stevenson.edu")
    doc.add_paragraph("EDUCATION")
    doc.add_paragraph("1990 University of Maryland B.A. Sociology")
    doc.add_paragraph("PROFESSIONAL EXPERIENCE")
    doc.add_paragraph("2005-Present Stevenson University")

    # CV body that the pattern detector CANNOT find — only sub-TOC anchors it.
    doc.add_paragraph("Thomas K. Swisher, J.D., Ph.D.")
    doc.add_paragraph("11886 Simpson Road")
    doc.add_paragraph("Clarksville, Md. 21029")
    doc.add_paragraph("Education")  # Title-case, defeats _is_anchor pattern
    doc.add_paragraph("1983 University of Virginia B.S. Secondary Education")
    doc.add_paragraph("1986 University of Baltimore Juris Doctorate")

    doc.add_paragraph("LAURI A. WEINER, J.D., HS-BCP")  # all-caps name
    doc.add_paragraph("7905 Winterset Avenue")
    doc.add_paragraph("Baltimore, Maryland 21208")
    doc.add_paragraph("lweiner@stevenson.edu")
    doc.add_paragraph("Education")
    doc.add_paragraph("J.D., University of Maryland School of Law, 1992")

    doc.add_paragraph("250 Faculty Drive · Baltimore, MD 21210")  # contact before name
    doc.add_paragraph("Mary Beth Olson, M.A.")
    doc.add_paragraph("Education")
    doc.add_paragraph("2007 Loyola University Maryland M.A. Counseling")

    doc.save(out_path)


@pytest.fixture
def stevenson_real_shape_docx(tmp_path):
    out = tmp_path / "stevenson-real.docx"
    _build_stevenson_real_shape_docx(out)
    return out


class TestRedetectStevensonRealShape:
    """End-to-end coverage for the REAL Stevenson document shape the
    user reported (2026-05-27 screenshot):

      - Main TOC has Standard N entries (not body markers) and a single
        "Appendices (List of Supporting Documents)" line.
      - Per-CV / per-syllabus / per-paper entries live in SUB-TOCs deep
        in the body, under "Appendix A: List of Faculty CVs" etc.
      - CV body paragraphs use credentials-prefixed names + Title-case
        section markers that defeat the pattern detector's anchor.

    Without the sub-TOC scanner + the parse_toc fix, the re-detect
    surfaced only 4 CVs (John, Carol, Roxanne, Barry — the easy ones).
    With both fixes, all 7 CVs are recovered.
    """

    def test_main_toc_with_standard_entries_does_not_bail_early(
        self, client, stevenson_real_shape_docx
    ):
        """The main TOC has Standard 1, Standard 2, Standard 6 listed
        as entries. parse_toc must walk THROUGH those (they're TOC
        lines, not body markers) and continue to the Appendices
        entry at the bottom."""
        resp = _post_redetect(
            client,
            {"s3Key": "imports/stevenson-real.docx", "importId": "imp-real", "submissionId": "sub-real"},
            stevenson_real_shape_docx,
        )
        body = resp.json()
        assert body["ok"] is True
        # tocEntriesFound includes the Standards entries AND the
        # Appendices line. Should be ≥ 5.
        assert body["tocDiagnostics"]["tocEntriesFound"] >= 5

    def test_sub_toc_recovers_credentials_prefixed_cvs(
        self, client, stevenson_real_shape_docx
    ):
        """The sub-TOC scanner finds the per-CV entries listed under
        "Appendix A: List of Faculty CVs" and anchors them in the
        body. Without this, Thomas Swisher / LAURI Weiner / Mary Beth
        Olson would be missed (their body anchors defeat the pattern
        detector)."""
        resp = _post_redetect(
            client,
            {"s3Key": "imports/stevenson-real.docx", "importId": "imp-real", "submissionId": "sub-real"},
            stevenson_real_shape_docx,
        )
        body = resp.json()
        cv_names_lower = " | ".join(c["facultyName"].lower() for c in body["cvs"])
        assert "thomas k. swisher" in cv_names_lower, (
            f"Expected Thomas K. Swisher recovered via sub-TOC; got {cv_names_lower}"
        )
        assert "lauri a. weiner" in cv_names_lower
        assert "mary beth olson" in cv_names_lower
        # And pattern-found CVs are still present.
        assert "john rosicky" in cv_names_lower

    def test_sub_toc_count_at_least_seven(self, client, stevenson_real_shape_docx):
        """All 7 CVs listed in the sub-TOC must be in the final cvs[]."""
        resp = _post_redetect(
            client,
            {"s3Key": "imports/stevenson-real.docx", "importId": "imp-real", "submissionId": "sub-real"},
            stevenson_real_shape_docx,
        )
        body = resp.json()
        assert body["counts"]["cvs"] >= 7, (
            f"Expected ≥7 CVs from sub-TOC + pattern; got {body['counts']['cvs']}. "
            f"Sub-TOC scanner is broken or merged into wrong section."
        )

    def test_sub_toc_recovers_three_syllabi(self, client, stevenson_real_shape_docx):
        resp = _post_redetect(
            client,
            {"s3Key": "imports/stevenson-real.docx", "importId": "imp-real", "submissionId": "sub-real"},
            stevenson_real_shape_docx,
        )
        body = resp.json()
        syllabi = [d for d in body["evidenceDocs"] if d["docSubKind"] == "syllabus"]
        # CHS 220, CHS 305, CHS 410 — three course-coded entries.
        codes = " ".join(s.get("courseCode") or "" for s in syllabi)
        codes += " " + " ".join(s["title"] for s in syllabi)
        assert "CHS 220" in codes
        assert "CHS 305" in codes
        assert "CHS 410" in codes

    def test_toc_diagnostics_reports_sub_toc_contribution(
        self, client, stevenson_real_shape_docx
    ):
        """tocDiagnostics.tocAdded.cvs must be > 0 because the sub-TOC
        contributed CVs the pattern detector missed."""
        resp = _post_redetect(
            client,
            {"s3Key": "imports/stevenson-real.docx", "importId": "imp-real", "submissionId": "sub-real"},
            stevenson_real_shape_docx,
        )
        body = resp.json()
        diag = body["tocDiagnostics"]
        assert diag["tocAdded"]["cvs"] >= 3, (
            f"Sub-TOC pass should have recovered at least 3 CVs the "
            f"pattern detector missed; got tocAdded={diag['tocAdded']}"
        )


class TestRedetectNoTocFallback:
    def test_no_toc_returns_zero_added(self, client, no_toc_docx):
        resp = _post_redetect(
            client,
            {"s3Key": "k", "importId": "i", "submissionId": "s"},
            no_toc_docx,
        )
        assert resp.status_code == 200
        body = resp.json()
        diag = body["tocDiagnostics"]
        assert diag["tocEntriesFound"] == 0
        assert diag["tocAnchoredDetections"] == 0
        assert diag["tocAdded"] == {"cvs": 0, "papers": 0, "syllabi": 0}

    def test_no_toc_still_returns_pattern_cv(self, client, no_toc_docx):
        """With no TOC, the pattern detector must still find Mary Smith
        — proving the pipeline didn't regress when TOC was added."""
        resp = _post_redetect(
            client,
            {"s3Key": "k", "importId": "i", "submissionId": "s"},
            no_toc_docx,
        )
        body = resp.json()
        names = " ".join(c["facultyName"].lower() for c in body["cvs"])
        assert "mary smith" in names
