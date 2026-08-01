"""Offline tests for the CSHSE Self-Study Template walker.

Covers:
  - canonical heading forms (``1.``, ``2a.``, ``11.b``, ``Standard 1, Specification a``)
  - ``Response:`` marker stripping
  - placeholder detection (Not applicable, TBD, See Appendix-only, empty body)
  - front-matter filtering (TOC, Introduction & Instructions)
  - body accumulation across multiple paragraphs
  - hint extraction (std + spec from heading)
"""
from __future__ import annotations

from app.splitter.template_walker import (
    walk_template_paragraphs,
    walk_template_docx,
    _is_placeholder,
    _match_heading,
)


def _texts(*paragraphs: str) -> list[str]:
    return list(paragraphs)


# ---------------------------------------------------------------- heading match


def test_match_heading_canonical_dotted_spec():
    assert _match_heading("1.a Provide…") == ("1", "a")
    assert _match_heading("11.b Describe…") == ("11", "b")


def test_match_heading_template_intro_form():
    assert _match_heading("2a. Describe the organizational structure") == ("2", "a")
    assert _match_heading("4c.  Describe any major program changes") == ("4", "c")


def test_match_heading_full_standard_specification_form():
    assert _match_heading("Standard 1, Specification a. The program is…") == ("1", "a")
    assert _match_heading("Standard 16 Specification b") == ("16", "b")
    assert _match_heading("Standard 12b. Small Groups") == ("12", "b")


def test_match_heading_standard_root_no_spec():
    assert _match_heading("1. Specify the degree(s)…") == ("1", None)
    assert _match_heading("Standard 1") == ("1", None)


def test_match_heading_section_letter_no_std():
    # "A. Required Self-Study Introductory Material" — section letter is
    # template-structural, not a Handbook standard hint.
    result = _match_heading("A. Required Self-Study Introductory Material")
    assert result is not None
    std, spec = result
    assert std is None  # no numeric std hint from a letter-prefix heading


def test_match_heading_rejects_body_paragraphs():
    assert _match_heading("KSU is governed by The Board of Regents…") is None
    assert _match_heading("Response:") is None
    assert _match_heading("See Appendix I") is None
    assert _match_heading("Not applicable") is None


def test_match_heading_skips_frontmatter():
    # Front-matter headings ARE cuts but emit no hints.
    assert _match_heading("Table of Contents") == (None, None)
    assert _match_heading("Introduction and Instructions: this template…") == (None, None)


# ---------------------------------------------------------------- placeholder


def test_placeholder_empty_body():
    assert _is_placeholder([]) is True
    assert _is_placeholder(["   ", ""]) is True


def test_placeholder_not_applicable():
    assert _is_placeholder(["Not applicable"]) is True
    assert _is_placeholder(["N/A"]) is True
    assert _is_placeholder(["TBD"]) is True


def test_placeholder_only_appendix_pointer():
    """A body that is JUST a 'See Appendix' line — no authored prose — is a placeholder.

    Reason: those pointers tell the institution where evidence will go,
    but until prose is written under the heading, the section is
    effectively unfilled.
    """
    assert _is_placeholder(["See Appendix I — #2 Map of Marietta Campus"]) is True


def test_not_placeholder_when_response_present():
    assert _is_placeholder(["The program was founded in 1999 by Dr. Franyo…"]) is False


def test_not_placeholder_when_appendix_pointer_plus_prose():
    body = [
        "See Appendix I — #1 Map of Kennesaw Campus",
        "KSU is a public research university with two campuses in Kennesaw and Marietta.",
    ]
    assert _is_placeholder(body) is False


# ---------------------------------------------------------------- walker


def test_walker_cuts_on_heading_and_accumulates_body():
    paras = _texts(
        "1. Specify the degree(s) offered",
        "Response:",
        "Bachelor of Science in Human Services.",
        "Accredited since 2004.",
        "2a. Describe the organizational structure",
        "Response:",
        "KSU is governed by the Board of Regents.",
    )
    sections = walk_template_paragraphs(paras)
    assert len(sections) == 2

    assert sections[0].heading.startswith("1. ")
    assert sections[0].standard_hint == "1"
    assert sections[0].spec_hint is None
    assert "Bachelor of Science in Human Services." in sections[0].body_text
    assert "Accredited since 2004." in sections[0].body_text
    # Response: marker should be stripped from the body
    assert "Response:" not in sections[0].body_text
    assert sections[0].placeholder is False

    assert sections[1].standard_hint == "2"
    assert sections[1].spec_hint == "a"
    assert "Board of Regents" in sections[1].body_text


def test_walker_marks_empty_section_as_placeholder():
    paras = _texts(
        "5a. Describe the physical location",
        "Response:",
        "",
        "5b. Describe the student population",
        "Response:",
        "The student population is 80% commuter.",
    )
    sections = walk_template_paragraphs(paras)
    assert len(sections) == 2
    assert sections[0].placeholder is True   # empty body
    assert sections[1].placeholder is False  # has real content


def test_walker_marks_not_applicable_as_placeholder():
    paras = _texts(
        "5. If the Program is delivered at multiple sites:",
        "Response:",
        "Not applicable",
        "6. Hybrid or Online Course Delivery",
        "Response:",
        "Not applicable",
    )
    sections = walk_template_paragraphs(paras)
    assert len(sections) == 2
    assert all(s.placeholder for s in sections)


def test_walker_skips_response_only_lines():
    """A bare ``Response:`` line emits nothing — actual content follows on the next paragraph."""
    paras = _texts(
        "1. Specify the degree(s) offered",
        "Response:",
        "Bachelor of Science in Human Services.",
    )
    sections = walk_template_paragraphs(paras)
    assert sections[0].body_text == "Bachelor of Science in Human Services."


def test_walker_strips_inline_response_marker():
    """``Response: text`` on one line should drop the marker but keep the text."""
    paras = _texts(
        "2a. Describe the organizational structure",
        "Response: KSU is governed by the Board of Regents.",
    )
    sections = walk_template_paragraphs(paras)
    assert sections[0].body_text == "KSU is governed by the Board of Regents."


def test_walker_drops_preamble_before_first_heading():
    """Paragraphs above the first heading (template title, instructions, etc.) are discarded."""
    paras = _texts(
        "Council for Standards in Human Service Education",
        "Self-Study Template",
        "BACCALAUREATE DEGREE IN HUMAN SERVICES",
        "Introduction and Instructions: This template is required…",
        "1. Specify the degree(s) offered",
        "Response:",
        "Bachelor of Science.",
    )
    sections = walk_template_paragraphs(paras)
    # Front-matter headings (Introduction and Instructions) ARE cuts but
    # have no body, so they become placeholder sections. We accept either
    # 1 or 2 depending on whether the front-matter section is filtered.
    # The important assertion: the substantive section is captured.
    real = [s for s in sections if not s.placeholder]
    assert len(real) == 1
    assert real[0].heading.startswith("1. ")
    assert real[0].body_text == "Bachelor of Science."


def test_walker_handles_per_spec_form_when_present():
    """Per-spec headings (``Standard 1, Specification a``) carry both std and spec hints."""
    paras = _texts(
        "Standard 1, Specification a.",
        "The program is part of Kennesaw State University, accredited by SACSCOC.",
        "Standard 1, Specification b.",
        "The primary objective of the program is to prepare HS professionals.",
    )
    sections = walk_template_paragraphs(paras)
    assert len(sections) == 2
    assert (sections[0].standard_hint, sections[0].spec_hint) == ("1", "a")
    assert (sections[1].standard_hint, sections[1].spec_hint) == ("1", "b")


def test_match_heading_standard_colon_form():
    """'Standard 1: The primary objective shall…' — colon form is a Standard
    ROOT (std hint, no spec). The shall-sentence runs on the same line."""
    assert _match_heading("Standard 1: The primary program objective shall be to prepare…") == ("1", None)
    assert _match_heading("Standard 10: Each program shall articulate policies…") == ("10", None)


def test_match_heading_widened_spec_letters():
    """Spec letters run past 'h' (e.g. Field Experience 21 a–j)."""
    assert _match_heading("21.j Provide…") == ("21", "j")
    assert _match_heading("9e. Describe office/classroom space") == ("9", "e")


# ------------------------------------------------ introduction region (KSU-style)


def _ksu_like():
    # Title block + intro prompts (their OWN 1./2a. numbering) + glossary,
    # THEN the standards region opening with "Standard 1: …" and real specs.
    return _texts(
        "1. Specify the degree(s) offered",
        "Response:",
        "Bachelor of Science in Human Services.",
        "2a. Describe the organizational structure",
        "Response:",
        "KSU is governed by the Board of Regents.",
        "B. Glossary of terms",
        "Accreditation: the process by which…",
        "Standard 1: The primary program objective shall be to prepare human services professionals.",
        "Context: institutional requirements.",
        "1a. The program is part of a regionally accredited institution.",
        "Response:",
        "KSU is accredited by SACSCOC.",
        "1b. Provide evidence that competent professionals is the primary objective.",
        "Response:",
        "The program's purpose is…",
    )


def test_intro_region_items_are_flagged_and_lose_spec_hints():
    sections = walk_template_paragraphs(_ksu_like())
    by_heading = {s.heading: s for s in sections}

    # The intro's own "1." / "2a." numbering must NOT be read as Standard specs.
    intro_degree = by_heading["1. Specify the degree(s) offered"]
    assert intro_degree.is_introduction is True
    assert (intro_degree.standard_hint, intro_degree.spec_hint) == (None, None)

    intro_org = by_heading["2a. Describe the organizational structure"]
    assert intro_org.is_introduction is True
    assert (intro_org.standard_hint, intro_org.spec_hint) == (None, None)

    # The glossary closes the introduction and stays in the intro region.
    glossary = by_heading["B. Glossary of terms"]
    assert glossary.is_introduction is True


def test_standards_after_intro_keep_their_hints():
    sections = walk_template_paragraphs(_ksu_like())
    by_heading = {s.heading: s for s in sections}

    # The Standard root (colon form) ends the intro and is NOT an intro section.
    std1 = next(s for s in sections if s.heading.startswith("Standard 1:"))
    assert std1.is_introduction is False
    assert std1.standard_hint == "1"
    assert std1.spec_hint is None

    # Real specs under the standard keep their std+spec hints.
    spec_1a = by_heading["1a. The program is part of a regionally accredited institution."]
    assert spec_1a.is_introduction is False
    assert (spec_1a.standard_hint, spec_1a.spec_hint) == ("1", "a")


def test_walker_keeps_see_appendix_pointer_inline():
    """``See Appendix`` lines mixed with prose are kept (matcher uses them as evidence pointers)."""
    paras = _texts(
        "2a. Describe the organizational structure",
        "Response:",
        "KSU is governed by the Board of Regents.",
        "See Appendix I — #1 Map of Kennesaw Campus",
        "The university operates on two campuses.",
    )
    sections = walk_template_paragraphs(paras)
    assert "See Appendix I — #1 Map of Kennesaw Campus" in sections[0].body_text
    assert "two campuses" in sections[0].body_text


# ------------------------------------- walk_template_docx flag emission (pipeline contract)


def _make_docx(paragraphs, path):
    from docx import Document as _D
    d = _D()
    for p in paragraphs:
        d.add_paragraph(p)
    d.save(str(path))


def test_docx_flags_route_intro_and_keep_spec_hints(tmp_path):
    """The flags the template pipeline reads: document intro + glossary →
    introduction:document; Standard root → introduction:standard-N; real spec
    keeps templateStandardHint/templateSpecHint and carries NO intro hint."""
    docx = tmp_path / "tmpl.docx"
    _make_docx(
        [
            "1. Specify the degree(s) offered",
            "Response:",
            "Bachelor of Science in Human Services.",
            "B. Glossary of terms",
            "Accreditation: the process by which an institution is reviewed.",
            "Standard 1: The primary program objective shall be to prepare professionals.",
            "Context: institutional requirements.",
            "1a. The program is part of a regionally accredited institution.",
            "Response:",
            "KSU is accredited by SACSCOC.",
        ],
        docx,
    )
    sections, _raw = walk_template_docx(str(docx), base_id="t")
    by_heading = {s.heading: s for s in sections}

    intro = by_heading["1. Specify the degree(s) offered"]
    assert intro.flags.get("templateIntroductionHint") == "introduction:document"
    assert intro.flags.get("templateStandardHint") == ""
    assert intro.flags.get("templateSpecHint") == ""

    glossary = by_heading["B. Glossary of terms"]
    assert glossary.flags.get("templateIntroductionHint") == "introduction:document"

    std_root = next(s for s in sections if s.heading.startswith("Standard 1:"))
    assert std_root.flags.get("templateIntroductionHint") == "introduction:standard-1"

    spec = by_heading["1a. The program is part of a regionally accredited institution."]
    assert spec.flags.get("templateIntroductionHint") in (None, "")
    assert spec.flags.get("templateStandardHint") == "1"
    assert spec.flags.get("templateSpecHint") == "a"


def test_inline_standard_citation_does_not_end_intro_before_glossary():
    """A front-matter curriculum citation ('Standard 13 a:') must NOT end the
    Introduction — the glossary that follows it still belongs to the intro."""
    paras = _texts(
        "1. Specify the degree(s) offered",
        "Response:",
        "Bachelor of Science.",
        "Standard 13 a: students understand small-group theory",  # inline citation
        "B. Include a glossary of terms",
        "Accreditation: the process by which…",
        "Standard 1: The primary objective shall be to prepare professionals.",
        "1a. The program is regionally accredited.",
        "Response:",
        "Accredited by SACSCOC.",
    )
    sections = walk_template_paragraphs(paras)
    glossary = next(s for s in sections if "glossary" in s.heading.lower())
    assert glossary.is_introduction is True
    spec_1a = next(s for s in sections if s.heading.startswith("1a."))
    assert spec_1a.is_introduction is False
    assert (spec_1a.standard_hint, spec_1a.spec_hint) == ("1", "a")


def test_standard_root_with_empty_body_still_emits_standard_intro(tmp_path):
    """Standard 9: with the shall-sentence in the heading and no body still
    emits a section routed to introduction:standard-9 (not dropped as empty)."""
    docx = tmp_path / "t.docx"
    _make_docx(
        [
            "I. GENERAL PROGRAM CHARACTERISTICS",
            "Standard 9: The program shall have adequate faculty, staff, and resources.",
            "9a. Include budgetary information.",
            "Response:",
            "The annual budget is sufficient.",
        ],
        docx,
    )
    sections, _ = walk_template_docx(str(docx), base_id="t")
    root = next(s for s in sections if s.heading.startswith("Standard 9:"))
    assert root.flags.get("templateIntroductionHint") == "introduction:standard-9"


def test_walker_captures_embedded_table_and_appendix_refs(tmp_path):
    """Embedded tables attach to the current spec as evidence; 'See Appendix'
    lines are captured as import reminders. (Paragraph-only walking dropped
    tables entirely — the rosters/grids that ARE the supporting evidence.)"""
    from docx import Document as _D
    d = _D()
    d.add_paragraph("I. GENERAL PROGRAM CHARACTERISTICS")
    d.add_paragraph("Standard 7: The program shall manage essential roles.")
    d.add_paragraph("7b 2. Provide a table matching faculty and staff with roles.")
    d.add_paragraph("Response:")
    d.add_paragraph("The roles are filled as shown. See Appendix VII Standard 7 Roster.")
    t = d.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "Name"; t.rows[0].cells[1].text = "Role"
    t.rows[1].cells[0].text = "Jane Doe"; t.rows[1].cells[1].text = "Director"
    docx = tmp_path / "roster.docx"; d.save(str(docx))

    sections, raw = walk_template_docx(str(docx), base_id="t")
    spec = next(r for r in raw if r.standard_hint == "7" and r.spec_hint == "b")
    assert len(spec.evidence_tables) == 1
    assert "Name" in spec.evidence_tables[0]["text"] and "Director" in spec.evidence_tables[0]["text"]
    assert "<table>" in spec.evidence_tables[0]["html"]
    assert spec.appendix_refs and "Appendix VII" in spec.appendix_refs[0]


def test_rich_html_preserves_links_images_tables(tmp_path):
    """Import fidelity: hyperlinks, inline images, and tables survive into the
    section's html_snippet (the plain-text walk dropped all three)."""
    import io, base64
    from docx import Document as _D
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn

    d = _D()
    d.add_paragraph("I. GENERAL PROGRAM CHARACTERISTICS")
    d.add_paragraph("Standard 1: The primary objective shall be to prepare professionals.")
    d.add_paragraph("1a. The program is regionally accredited.")
    d.add_paragraph("Response:")
    p = d.add_paragraph("Accredited by ")
    rid = p.part.relate_to(
        "https://sacscoc.org/",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hl = OxmlElement("w:hyperlink"); hl.set(_qn("r:id"), rid)
    r = OxmlElement("w:r"); t = OxmlElement("w:t"); t.text = "SACSCOC"; r.append(t); hl.append(r)
    p._p.append(hl)
    # image-only paragraph — a valid generated 1x1 PNG
    import struct, zlib
    def _png_1x1():
        def _c(typ, data):
            body = typ + data
            return struct.pack(">I", len(data)) + body + struct.pack(">I", zlib.crc32(body) & 0xffffffff)
        ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
        idat = zlib.compress(b"\x00\xff\x00\x00")
        return b"\x89PNG\r\n\x1a\n" + _c(b"IHDR", ihdr) + _c(b"IDAT", idat) + _c(b"IEND", b"")
    d.add_picture(io.BytesIO(_png_1x1()))
    tb = d.add_table(rows=1, cols=2); tb.rows[0].cells[0].text = "Name"; tb.rows[0].cells[1].text = "Role"
    docx = tmp_path / "rich.docx"; d.save(str(docx))

    sections, _ = walk_template_docx(str(docx), base_id="t")
    spec = next(s for s in sections if s.heading.startswith("1a."))
    html = spec.html_snippet or ""
    assert 'href="https://sacscoc.org/"' in html and "SACSCOC" in html   # hyperlink kept
    assert '<img src="data:image/png;base64,' in html                    # inline image kept
    assert "<table>" in html and "Name" in html                          # table inline
    assert spec.contains_image is True


def test_cr061_consolidates_subitems_into_spec_until_next_break():
    """CR-061 — 'write everything until the next break': a spec owns ALL content
    (sub-numbered items "1b 1.", bare list numbers "1."/"2.") until the NEXT
    spec/standard heading, not until the next numbered line. Previously these
    fragmented one spec across many sections (then mis-matched / dropped)."""
    paras = [
        "I. GENERAL PROGRAM CHARACTERISTICS",
        "Standard 1: The primary objective shall be X.",
        "1a. The program is part of a degree granting college.",
        "Response:",
        "KSU is accredited by SACSCOC.",
        "1b. Provide evidence of the primary objective.",
        "The HS program prepares students.",
        "1b 1. A detailed description of the membership.",
        "Members include faculty and staff.",
        "1. The board met on December 1.",
        "Minutes attached.",
        "2. The committee met on March 1.",
        "1c. Articulate how students are informed.",
        "Students are informed via the catalog.",
    ]
    sections = walk_template_paragraphs(paras)
    specs = {
        (s.standard_hint, s.spec_hint): s
        for s in sections
        if s.standard_hint and s.spec_hint
    }
    # exactly one section per spec — no fragmentation
    assert ("1", "a") in specs and ("1", "b") in specs and ("1", "c") in specs
    b = specs[("1", "b")].body_text
    # 1b owns everything until 1c: its own text + every sub-item
    assert "prepares students" in b
    assert "detailed description of the membership" in b   # "1b 1." sub-item kept
    assert "board met on December 1" in b                  # bare "1." sub-item kept
    assert "committee met on March 1" in b                 # bare "2." sub-item kept
    # but NOT 1c's content (1c is the next break)
    assert "informed via the catalog" not in b
    assert "informed via the catalog" in specs[("1", "c")].body_text


def test_cr062_consecutive_li_wrap_in_ul():
    """CR-062 — consecutive <li> items wrap in one <ul>; links inside survive;
    surrounding <p> paragraphs are untouched."""
    from app.splitter.template_walker import TemplateSection
    s = TemplateSection(
        paragraph_index=0, heading="1a.", body_paragraphs=[],
        standard_hint="1", spec_hint="a", placeholder=False,
    )
    s.body_html_parts = [
        "<p>Intro</p>",
        '<li><a href="https://x.edu">KSU Catalog</a></li>',
        "<li>Student Handbook</li>",
        "<p>After</p>",
    ]
    html = s.body_html
    assert '<ul><li><a href="https://x.edu">KSU Catalog</a></li><li>Student Handbook</li></ul>' in html
    assert "<p>Intro</p>" in html and "<p>After</p>" in html
    assert html.count("<ul>") == 1


def test_cr062_is_list_item_detects_numpr():
    """CR-062 — _is_list_item flags Word numbered/bulleted paragraphs."""
    import docx
    from docx.oxml.ns import qn
    from app.splitter.template_walker import _is_list_item
    d = docx.Document()
    p = d.add_paragraph("bullet")
    pPr = p._p.get_or_add_pPr()
    pPr.append(pPr.makeelement(qn("w:numPr"), {}))
    assert _is_list_item(p._p) is True
    assert _is_list_item(d.add_paragraph("plain")._p) is False


def test_cr063_image_cap_raised_to_2mb():
    """CR-063 — inline-image cap raised so realistic images aren't dropped."""
    from app.splitter import template_walker
    assert template_walker._MAX_INLINE_IMAGE_BYTES >= 2_000_000


def test_standard_context_descriptor_does_not_bleed_into_prior_spec():
    """A standard's TITLE + 'Context:' rubric block must NOT accumulate onto the
    previous spec's response (the AACC boundary-bleed bug). Prod excludes them."""
    from app.splitter.template_walker import walk_template_paragraphs
    paras = [
        "Standard 4: The program shall document achievement.",
        "4.c. Provide data.",
        "Response:",
        "The program publishes achievement data on its website.",
        # --- next standard's descriptor block (title + Context:) ---
        "Policies and Procedures for Admitting, Retaining, and Dismissing Students",
        "Context: Students have a right to know, prior to enrollment, the standards "
        "of the program and the procedures for admitting, retaining, and dismissing students.",
        "Standard 5: The program shall have written standards and procedures.",
        "5.a. Provide documentation of admission policies.",
        "Response:",
        "The program admits students per the catalog.",
    ]
    secs = walk_template_paragraphs(paras)
    joined = " ".join(s.body_text for s in secs)
    # the 4.c response is kept, the descriptor block is dropped
    assert "publishes achievement data" in joined
    assert "admits students per the catalog" in joined
    assert "Context:" not in joined
    assert "admitting, retaining, and dismissing students" not in joined
    # the standard-5 title line did not bleed into 4.c
    four_c = next((s for s in secs if s.spec_hint == "c" and s.standard_hint == "4"), None)
    assert four_c is not None
    assert "Policies and Procedures" not in four_c.body_text
